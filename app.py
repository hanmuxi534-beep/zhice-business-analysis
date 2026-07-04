# -*- coding: utf-8 -*-
"""
智策经营——AI 驱动的多维经营分析与决策支持系统
高级视觉版

重点：
1. 不再使用旧版 plotly express 的高风险函数，避免 pandas/plotly 兼容问题；
2. 数值列采用稳健解析，保留 Excel 中真实数值；
3. 图表绘制前先检查有效数据，不再出现空白图却没有说明的情况；
4. 有日期看趋势，无日期看结构；
5. 异常诊断解释“为什么异常”和“为什么高风险”；
6. 问数示例根据当前上传数据动态生成。
"""

import os
import re
import time
import json
import sqlite3
from io import BytesIO

import numpy as np
import pandas as pd
import requests
import streamlit as st
import plotly.graph_objects as go
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from sklearn.ensemble import IsolationForest
from sklearn.preprocessing import StandardScaler

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 0. 基础设置
# ============================================================

st.set_page_config(
    page_title="智策经营",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
:root {
    --zc-primary: #123A63;
    --zc-primary-2: #1E77D3;
    --zc-bg: #F3F6FB;
    --zc-card: #FFFFFF;
    --zc-text: #10213F;
    --zc-muted: #66758C;
    --zc-line: #E5EBF3;
    --zc-shadow: 0 14px 34px rgba(24, 48, 88, .08);
}

.stApp { background: radial-gradient(circle at top left, #F8FBFF 0, #F3F6FB 42%, #EEF3FA 100%); }
.block-container { padding-top: 1.1rem; padding-bottom: 2.6rem; max-width: 1680px; }
html, body, [class*="css"] { font-size: 17px; }
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #EEF4FB 0%, #E8EEF6 100%);
    border-right: 1px solid #DDE6F1;
    min-width: 350px !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: var(--zc-text);
}
section[data-testid="stSidebar"] div[data-testid="stMarkdownContainer"] p,
section[data-testid="stSidebar"] li {
    font-size: 16px;
    line-height: 1.7;
}
label[data-testid="stWidgetLabel"] {
    font-size: 16px !important;
    font-weight: 800 !important;
    color: #243654 !important;
}
.stSelectbox div[data-baseweb="select"],
.stMultiSelect div[data-baseweb="select"],
.stTextInput input {
    min-height: 48px;
    border-radius: 14px !important;
    font-size: 16px !important;
}
.stButton button {
    border-radius: 14px !important;
    min-height: 46px;
    font-size: 17px !important;
    font-weight: 800 !important;
    padding: 0 24px !important;
}
.stDownloadButton button {
    border-radius: 14px !important;
    min-height: 46px;
    font-size: 17px !important;
    font-weight: 800 !important;
}
.stTabs [data-baseweb="tab-list"] {
    gap: 10px;
    border-bottom: 1px solid #DDE6F1;
    padding-bottom: 6px;
}
.stTabs [data-baseweb="tab"] {
    font-size: 18px !important;
    font-weight: 900 !important;
    padding: 14px 20px !important;
    border-radius: 14px 14px 0 0;
}
.stTabs [aria-selected="true"] {
    background: #FFFFFF !important;
    box-shadow: 0 -2px 14px rgba(24, 48, 88, .05);
}
h1 { font-size: 40px !important; font-weight: 950 !important; color: var(--zc-text); }
h2 { font-size: 34px !important; font-weight: 950 !important; color: var(--zc-text); margin-top: 18px !important; }
h3 { font-size: 25px !important; font-weight: 900 !important; color: var(--zc-text); margin-top: 16px !important; }
div[data-testid="stMarkdownContainer"] p,
div[data-testid="stMarkdownContainer"] li {
    font-size: 17px;
    line-height: 1.85;
}
.hero {
    background:
      linear-gradient(135deg, rgba(16, 52, 92, .98), rgba(30, 119, 211, .94)),
      radial-gradient(circle at 92% 8%, rgba(255,255,255,.22), transparent 26%);
    border-radius: 28px;
    padding: 34px 38px;
    color: white;
    box-shadow: 0 22px 42px rgba(15, 61, 99, .20);
    margin-bottom: 24px;
    border: 1px solid rgba(255,255,255,.22);
}
.hero h1 { font-size: 42px !important; line-height: 1.22; margin: 0; font-weight: 950; color: white !important; }
.hero p { font-size: 18px; opacity: .96; margin-top: 12px; max-width: 980px; }
.metric-card {
    background: linear-gradient(180deg, #FFFFFF 0%, #FBFDFF 100%);
    border: 1px solid var(--zc-line);
    border-radius: 24px;
    box-shadow: var(--zc-shadow);
    padding: 24px 26px;
    min-height: 158px;
    position: relative;
    overflow: hidden;
}
.metric-card:before {
    content: "";
    position: absolute;
    top: 0;
    left: 0;
    height: 5px;
    width: 100%;
    background: linear-gradient(90deg, #1E77D3, #78B7FF);
}
.metric-title { color: var(--zc-muted); font-size: 16px; font-weight: 800; margin-bottom: 12px; }
.metric-value { color: #071F43; font-size: 38px; font-weight: 950; letter-spacing: .2px; line-height: 1.1; }
.metric-note { color: #8290A5; font-size: 14px; margin-top: 12px; line-height: 1.55; }
.section-card {
    background: var(--zc-card);
    border: 1px solid var(--zc-line);
    border-radius: 24px;
    box-shadow: var(--zc-shadow);
    padding: 26px 28px;
    margin-bottom: 20px;
}
.tip-card {
    background: linear-gradient(90deg, #EAF4FF 0%, #F6FBFF 100%);
    border-left: 7px solid #2F80ED;
    padding: 18px 22px;
    border-radius: 18px;
    color: #0B3A6A;
    margin-bottom: 18px;
    line-height: 1.85;
    font-size: 17px;
    box-shadow: 0 8px 18px rgba(47, 128, 237, .08);
}
.warn-card {
    background: linear-gradient(90deg, #FFF7E1 0%, #FFFDF5 100%);
    border-left: 7px solid #F2A900;
    padding: 18px 22px;
    border-radius: 18px;
    color: #704E00;
    margin-bottom: 18px;
    line-height: 1.85;
    font-size: 17px;
    box-shadow: 0 8px 18px rgba(242, 169, 0, .10);
}
.risk-card {
    background: #FFFFFF;
    border: 1px solid var(--zc-line);
    border-radius: 22px;
    padding: 20px 22px;
    min-height: 190px;
    box-shadow: var(--zc-shadow);
}
.example-card {
    background: #FFFFFF;
    border: 1px solid var(--zc-line);
    border-radius: 18px;
    padding: 17px 18px;
    min-height: 118px;
    box-shadow: 0 8px 20px rgba(30,55,90,.06);
}
.example-title { font-size: 15px; color: #60718A; font-weight: 900; margin-bottom: 8px; }
.example-q { font-size: 17px; color: #0B2142; line-height: 1.65; font-weight: 650; }
.small-text { color: #5E6F86; font-size: 16px; line-height: 1.85; }
[data-testid="stDataFrame"] {
    border-radius: 18px;
    overflow: hidden;
    border: 1px solid #E3EAF4;
    box-shadow: 0 8px 20px rgba(30,55,90,.05);
}
.js-plotly-plot {
    border-radius: 22px;
    overflow: hidden;
    border: 1px solid #E5EBF3;
    box-shadow: 0 10px 24px rgba(30,55,90,.06);
    background: #FFFFFF;
}
hr { border-color: #DDE6F1; }

.field-card {
    background: linear-gradient(180deg, #FFFFFF 0%, #FBFDFF 100%);
    border: 1px solid #E5EBF3;
    border-radius: 22px;
    padding: 20px 22px;
    min-height: 178px;
    box-shadow: 0 10px 24px rgba(30,55,90,.06);
}
.field-card-title {
    font-size: 19px;
    font-weight: 950;
    color: #10213F;
    margin-bottom: 10px;
}
.field-card-desc {
    font-size: 15px;
    color: #66758C;
    line-height: 1.65;
    margin-bottom: 12px;
}
.pill-wrap { display: flex; flex-wrap: wrap; gap: 8px; }
.field-pill {
    display: inline-block;
    background: #EEF5FF;
    color: #164B80;
    border: 1px solid #D6E8FF;
    border-radius: 999px;
    padding: 7px 12px;
    font-size: 15px;
    font-weight: 800;
}
.field-pill.warn {
    background: #FFF7E6;
    color: #8A5A00;
    border-color: #FFE2A8;
}
.field-pill.muted {
    background: #F2F5F9;
    color: #66758C;
    border-color: #E1E7F0;
}
.health-summary {
    background: linear-gradient(135deg, #FFFFFF 0%, #F7FBFF 100%);
    border: 1px solid #E5EBF3;
    border-radius: 24px;
    padding: 24px 28px;
    box-shadow: 0 12px 28px rgba(30,55,90,.07);
    min-height: 300px;
}
.health-summary h3 {
    margin-top: 0 !important;
    margin-bottom: 12px !important;
}
.health-line {
    font-size: 17px;
    line-height: 1.85;
    color: #40516A;
}
.progress-row {
    margin: 12px 0 16px 0;
}
.progress-label {
    display: flex;
    justify-content: space-between;
    font-size: 15px;
    font-weight: 850;
    color: #243654;
    margin-bottom: 6px;
}
.progress-track {
    width: 100%;
    height: 14px;
    background: #E9EEF6;
    border-radius: 999px;
    overflow: hidden;
}
.progress-fill {
    height: 14px;
    border-radius: 999px;
    background: linear-gradient(90deg, #1E77D3, #71B7FF);
}
.detail-expander-note {
    color: #66758C;
    font-size: 15px;
    line-height: 1.7;
    margin-bottom: 10px;
}


/* 汇报增强补充 */
div[data-testid="stMarkdownContainer"] p,
div[data-testid="stMarkdownContainer"] li {
    font-size: 18px !important;
    line-height: 1.95 !important;
}
.small-text { font-size: 18px !important; line-height: 1.95 !important; }
.detail-expander-note { font-size: 17px !important; line-height: 1.85 !important; }
.highlight-note {
    background: linear-gradient(90deg, #F4F8FF 0%, #FFFFFF 100%);
    border-left: 6px solid #1E77D3;
    border-radius: 16px;
    padding: 15px 18px;
    font-size: 18px;
    color: #334766;
    line-height: 1.9;
    margin: 12px 0 16px 0;
}
.ai-table-card {
    background: #FFFFFF;
    border: 1px solid #E5EBF3;
    border-radius: 18px;
    padding: 18px 20px;
    box-shadow: 0 8px 22px rgba(30,55,90,.06);
    margin: 12px 0 18px 0;
}
.ai-section-title {
    font-size: 21px;
    font-weight: 950;
    color: #10213F;
    margin: 10px 0 8px 0;
}


.asset-grid {
    display: grid;
    grid-template-columns: repeat(5, minmax(140px, 1fr));
    gap: 16px;
    margin: 14px 0 20px 0;
}
.asset-card {
    background: linear-gradient(180deg, #FFFFFF 0%, #F9FCFF 100%);
    border: 1px solid #E3EBF6;
    border-radius: 22px;
    padding: 18px 20px;
    box-shadow: 0 10px 24px rgba(30,55,90,.065);
    position: relative;
    overflow: hidden;
    min-height: 122px;
}
.asset-card:before {
    content: "";
    position: absolute;
    left: 0;
    top: 0;
    width: 5px;
    height: 100%;
    background: linear-gradient(180deg, #1E77D3, #7BC8FF);
}
.asset-label {
    color: #66758C;
    font-size: 15px;
    font-weight: 900;
    margin-bottom: 8px;
}
.asset-value {
    color: #071F43;
    font-size: 31px;
    font-weight: 950;
    line-height: 1.15;
}
.asset-note {
    color: #7B8BA2;
    font-size: 14px;
    margin-top: 9px;
    line-height: 1.55;
}
.governance-hero {
    background:
      linear-gradient(135deg, rgba(255,255,255,.96) 0%, rgba(245,250,255,.98) 100%);
    border: 1px solid #DDEBFA;
    border-radius: 26px;
    padding: 24px 28px;
    box-shadow: 0 14px 34px rgba(24,48,88,.08);
    margin-bottom: 18px;
}
.governance-hero-title {
    font-size: 24px;
    font-weight: 950;
    color: #10213F;
    margin-bottom: 8px;
}
.governance-hero-text {
    font-size: 17px;
    color: #40516A;
    line-height: 1.85;
}
.governance-step {
    display: inline-block;
    margin: 6px 8px 6px 0;
    padding: 8px 13px;
    border-radius: 999px;
    background: #EEF6FF;
    color: #13528A;
    border: 1px solid #D7E9FF;
    font-size: 14px;
    font-weight: 900;
}
.action-list-card {
    background: #FFFFFF;
    border: 1px solid #E5EBF3;
    border-radius: 22px;
    padding: 20px 22px;
    box-shadow: 0 10px 24px rgba(30,55,90,.06);
    min-height: 300px;
}
.action-item {
    display: flex;
    justify-content: space-between;
    align-items: flex-start;
    gap: 14px;
    border-bottom: 1px solid #EDF2F8;
    padding: 12px 0;
}
.action-item:last-child { border-bottom: none; }
.action-title {
    font-weight: 900;
    color: #10213F;
    font-size: 16px;
}
.action-desc {
    color: #65768E;
    font-size: 14px;
    line-height: 1.6;
    margin-top: 4px;
}
.action-badge {
    white-space: nowrap;
    border-radius: 999px;
    padding: 5px 10px;
    font-size: 13px;
    font-weight: 900;
    background: #EAF5FF;
    color: #1E77D3;
}
.action-badge.warn {
    background: #FFF2D8;
    color: #A76700;
}
.action-badge.ok {
    background: #EAF8F0;
    color: #16834A;
}


.gov-layout {
    display: grid;
    grid-template-columns: 0.86fr 1.45fr;
    gap: 18px;
    margin-bottom: 18px;
}
.asset-grid-compact {
    display: grid;
    grid-template-columns: repeat(2, minmax(150px, 1fr));
    gap: 14px;
}
.asset-card.compact {
    min-height: 118px;
    padding: 18px 20px;
}
.compact-progress-wrap {
    background: #FFFFFF;
    border: 1px solid #E5EBF3;
    border-radius: 22px;
    padding: 18px 22px;
    box-shadow: 0 10px 24px rgba(30,55,90,.06);
    margin-bottom: 18px;
}
.compact-progress-row {
    display: grid;
    grid-template-columns: 150px 1fr 70px;
    align-items: center;
    gap: 14px;
    padding: 8px 0;
    border-bottom: 1px solid #EEF3F8;
}
.compact-progress-row:last-child { border-bottom: none; }
.compact-progress-name {
    color: #243654;
    font-weight: 900;
    font-size: 15px;
}
.compact-progress-track {
    height: 12px;
    background: #E9EEF6;
    border-radius: 999px;
    overflow: hidden;
}
.compact-progress-fill {
    height: 12px;
    border-radius: 999px;
    background: linear-gradient(90deg, #1E77D3, #71B7FF);
}
.compact-progress-value {
    color: #66758C;
    font-size: 14px;
    font-weight: 900;
    text-align: right;
}
.gov-section-title {
    font-size: 22px;
    font-weight: 950;
    color: #10213F;
    margin: 18px 0 10px 0;
}
@media (max-width: 1100px) {
    .gov-layout { grid-template-columns: 1fr; }
    .asset-grid-compact { grid-template-columns: 1fr; }
}

</style>
""", unsafe_allow_html=True)


# ============================================================
# 1. 数据读取与字段识别
# ============================================================

def normalize_column_name(col):
    col = str(col).strip()
    col = re.sub(r"\s+", "_", col)
    col = re.sub(r"[^\w\u4e00-\u9fff]", "_", col)
    col = re.sub(r"_+", "_", col)
    return col.strip("_")


def read_uploaded_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        try:
            return pd.read_csv(uploaded_file, encoding="utf-8")
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding="gbk")
    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file)
    raise ValueError("暂不支持该文件类型，请上传 xlsx、xls 或 csv。")


def is_id_like(col):
    c = str(col).lower()
    patterns = [
        r"(^|_)id($|_)", r"编号", r"序号", r"代码", r"code", r"key",
        r"index", r"row", r"流水", r"唯一", r"主键", r"no$"
    ]
    return any(re.search(p, c) for p in patterns)


def is_rate_like(col):
    c = str(col).lower()
    return any(k in c for k in ["rate", "ratio", "percent", "percentage", "pct", "margin", "discount", "折扣", "率", "占比", "比例", "完成率"])


def is_amount_like(col):
    c = str(col).lower()
    return any(k in c for k in [
        "amount", "cost", "expense", "spend", "revenue", "sales", "profit",
        "pay", "price", "fee", "total", "asset", "liab", "equity", "cash",
        "收入", "销售", "成本", "费用", "支出", "利润", "薪酬", "工资", "金额", "单价", "总额", "预算", "资产", "负债"
    ])


def is_quantity_like(col):
    c = str(col).lower()
    return any(k in c for k in ["qty", "quantity", "count", "number", "hours", "usage", "数量", "工时", "用量", "次数", "人数"])


def is_date_name_like(col):
    c = str(col).lower()
    return any(k in c for k in ["date", "time", "month", "year", "period", "日期", "时间", "月份", "年月", "期间"])


def is_progress_like(col):
    c = str(col).lower()
    return any(k in c for k in ["day_of", "days", "duration", "进度", "天数", "周期", "耗时", "时长"])



def aggregation_method_for_metric(col):
    """
    经营指标聚合口径：
    - 金额、收入、成本、利润、数量等规模型指标：求和；
    - 折扣、比率、百分比、单价、均价、时长、周期等水平型指标：平均；
    这样可以避免“折扣被相加成 3.05 / 4.40”这类明显不合理的问题。
    """
    c = str(col).lower()
    if is_rate_like(col):
        return "mean"
    if any(k in c for k in ["price", "unit_price", "avg", "average", "单价", "均价", "平均"]):
        return "mean"
    if is_progress_like(col):
        return "mean"
    return "sum"


def aggregation_label_for_metric(col):
    method = aggregation_method_for_metric(col)
    return "平均值" if method == "mean" else "合计"




def parse_numeric_series(s, aggressive=False):
    """稳健数值解析，支持 Excel 数值、逗号、货币符号、万/亿、百分号、Day 12 等。"""
    if pd.api.types.is_numeric_dtype(s):
        return pd.to_numeric(s, errors="coerce")

    raw = s.astype(str).str.strip()
    raw = raw.replace({
        "": np.nan, "nan": np.nan, "NaN": np.nan, "None": np.nan,
        "null": np.nan, "NULL": np.nan, "-": np.nan, "—": np.nan
    })

    neg_mask = raw.str.match(r"^\(.*\)$", na=False)
    cleaned = raw.str.replace(r"^\((.*)\)$", r"\1", regex=True)
    cleaned = cleaned.str.replace(r"[,，\s]", "", regex=True)
    cleaned = cleaned.str.replace(r"(人民币|RMB|CNY|USD|HKD|￥|¥|\$|元)", "", regex=True)

    percent_mask = cleaned.str.contains(r"%|％", na=False)

    multiplier = pd.Series(1.0, index=s.index)
    multiplier = multiplier.mask(cleaned.str.contains("亿", na=False), 100000000.0)
    multiplier = multiplier.mask(cleaned.str.contains("万", na=False), 10000.0)
    multiplier = multiplier.mask(cleaned.str.contains(r"[kK]$", na=False), 1000.0)
    multiplier = multiplier.mask(cleaned.str.contains(r"[mM]$", na=False), 1000000.0)

    cleaned = cleaned.str.replace(r"(亿|万|%|％|[kKmM])", "", regex=True)
    parsed = pd.to_numeric(cleaned, errors="coerce")

    if aggressive:
        extracted = cleaned.str.extract(r"([-+]?\d*\.?\d+(?:[eE][-+]?\d+)?)")[0]
        extracted = pd.to_numeric(extracted, errors="coerce")
        parsed = parsed.fillna(extracted)

    parsed = parsed * multiplier
    parsed = parsed.mask(percent_mask, parsed / 100.0)
    parsed = parsed.mask(neg_mask, -parsed)
    return parsed


def try_parse_date(series):
    parsed = pd.to_datetime(series, errors="coerce")
    valid_ratio = parsed.notna().mean()
    if valid_ratio == 0:
        return parsed, 0, False, "无法解析为日期"

    years = parsed.dropna().dt.year
    unique_dates = parsed.dropna().nunique()
    suspicious_epoch = (years.between(1969, 1971).mean() > 0.8) if len(years) else False
    too_few_unique = unique_dates < 2
    valid = valid_ratio >= 0.55 and not suspicious_epoch and not too_few_unique

    if valid_ratio < 0.55:
        reason = "日期解析比例较低"
    elif suspicious_epoch:
        reason = "疑似数值被误解析为1970附近日期"
    elif too_few_unique:
        reason = "有效日期取值过少"
    else:
        reason = "有效日期字段"
    return parsed, valid_ratio, valid, reason


def infer_field_metadata(raw_df):
    rows = []
    for col in raw_df.columns:
        s = raw_df[col]
        sample_values = s.dropna().astype(str).head(3).tolist()
        numeric = parse_numeric_series(s, aggressive=False)
        nr = numeric.notna().mean()
        uniq = int(s.nunique(dropna=True))
        ur = uniq / max(len(s), 1)

        if is_date_name_like(col):
            _, dr, valid_date, date_reason = try_parse_date(s)
        else:
            dr, valid_date, date_reason = 0, False, ""

        if is_id_like(col):
            role, ftype, agg, ok = "标识字段", "ID/编号", "去重计数 / 不聚合", "否"
            meaning = "用于唯一标识记录、对象或业务单元，主要用于追踪和去重，不应作为可加总的经营指标。"
            note = "不建议用于相关性分析或趋势分析。"
        elif valid_date:
            role, ftype, agg, ok = "时间字段", "日期", "按年/月/日分组", "否"
            meaning = "表示业务发生或统计所属的时间，可用于趋势、环比、同比和周期变化分析。"
            note = date_reason
        elif nr > 0.75 and is_progress_like(col):
            role, ftype, agg, ok = "进度/时长字段", "数值", "均值 / 分段分析", "可作为辅助指标"
            meaning = "表示项目进行天数、周期、进度或耗时，适合做进度型分析，不应直接当作自然日期。"
            note = "可作为数值横轴或辅助指标。"
        elif nr > 0.75 and is_rate_like(col):
            role, ftype, agg, ok = "度量指标", "比例/比率", "平均 / 加权平均", "谨慎"
            meaning = "表示比例、比率、折扣或百分比类指标，适合比较水平，不适合直接求和。"
            note = "默认建议使用平均或加权平均；折扣字段不能求和。"
        elif nr > 0.75 and is_amount_like(col):
            role, ftype, agg, ok = "度量指标", "金额/成本/收入类", "求和 / 均值", "是"
            meaning = "表示金额、成本、费用、收入、利润或预算等经营度量，可作为主分析指标。"
            note = "适合做总量、均值、排名、结构和异常分析。"
        elif nr > 0.75 and is_quantity_like(col):
            role, ftype, agg, ok = "度量指标", "数量/工时/用量类", "求和 / 均值", "是"
            meaning = "表示数量、工时、用量或次数等投入/产出规模指标，可用于规模、效率或驱动关系分析。"
            note = "可与金额类指标做关系分析。"
        elif nr > 0.75:
            role, ftype, agg, ok = "数值字段", "数值", "求和 / 均值 / 中位数", "可选"
            meaning = "数值型字段，可根据业务含义作为主指标或辅助指标分析。"
            note = "建议结合字段含义确认是否可加总。"
        elif ur < 0.65:
            role, ftype, agg, ok = "维度字段", "类别/分组", "分组统计", "否"
            meaning = "表示类别、部门、地区、状态、类型或对象分组，可作为经营洞察口径。"
            note = "适合用于分组对比、贡献度分析和结构分析。"
        else:
            role, ftype, agg, ok = "文本字段", "文本", "不聚合 / 计数", "否"
            meaning = "文本描述类字段，通常用于展示或辅助检索，不适合作为数值分析指标。"
            note = "如需分析文本内容，可后续接入文本挖掘。"

        rows.append({
            "字段名": col,
            "字段含义解释": meaning,
            "字段类型": ftype,
            "推荐角色": role,
            "推荐聚合方式": agg,
            "是否适合作为主指标": ok,
            "唯一值数量": uniq,
            "数值解析比例": round(float(nr), 3),
            "日期解析比例": round(float(dr), 3),
            "注意事项": note,
            "示例值": "、".join(sample_values)
        })
    return pd.DataFrame(rows)


def get_metric_candidates(meta):
    temp = meta[meta["推荐角色"].isin(["度量指标", "数值字段", "进度/时长字段"])]
    return temp["字段名"].tolist()


def get_dimension_candidates(meta):
    return meta[meta["推荐角色"].isin(["维度字段"])]["字段名"].tolist()


def get_date_candidates(meta):
    return meta[meta["推荐角色"].isin(["时间字段"])]["字段名"].tolist()


def convert_selected_numeric(df, cols):
    out = df.copy()
    report = []
    for col in cols:
        if col in out.columns:
            before_valid = out[col].notna().sum()
            out[col] = parse_numeric_series(out[col], aggressive=True)
            after_valid = out[col].notna().sum()
            report.append({
                "字段": col,
                "转换前非空数": int(before_valid),
                "转换后有效数值数": int(after_valid),
                "缺失/无法解析数": int(out[col].isna().sum())
            })
    return out, pd.DataFrame(report)


def clean_data(raw_df, selected_numeric_cols, date_col, missing_strategy):
    clean = raw_df.copy()
    raw_rows = len(clean)
    dup = int(clean.duplicated().sum())
    clean = clean.drop_duplicates()

    reports = [{"处理环节": "删除完全重复行", "处理数量": dup, "说明": "整行完全一致时视为重复记录。"}]

    if date_col and date_col in clean.columns:
        parsed, ratio, valid, reason = try_parse_date(clean[date_col])
        clean[date_col] = parsed
        reports.append({"处理环节": "日期字段解析", "处理数量": int(parsed.isna().sum()), "说明": f"{date_col}：{reason}，解析成功比例 {ratio:.1%}。"})

    clean, numeric_report = convert_selected_numeric(clean, selected_numeric_cols)
    for _, r in numeric_report.iterrows():
        reports.append({"处理环节": f"数值转换：{r['字段']}", "处理数量": int(r["缺失/无法解析数"]), "说明": f"转换后有效数值 {int(r['转换后有效数值数'])} 条。"})

    if missing_strategy == "删除主指标缺失行" and selected_numeric_cols:
        main = selected_numeric_cols[0]
        before = len(clean)
        clean = clean.dropna(subset=[main])
        reports.append({"处理环节": "缺失值处理", "处理数量": before - len(clean), "说明": f"删除主指标 {main} 缺失的记录。"})
    elif missing_strategy == "数值中位数填充，类别填未知":
        fill_count = 0
        for col in clean.columns:
            na = int(clean[col].isna().sum())
            if na == 0:
                continue
            if pd.api.types.is_numeric_dtype(clean[col]):
                clean[col] = clean[col].fillna(clean[col].median())
            else:
                clean[col] = clean[col].fillna("未知")
            fill_count += na
        reports.append({"处理环节": "缺失值处理", "处理数量": fill_count, "说明": "数值列用中位数填充，类别/文本列填未知。"})
    else:
        reports.append({"处理环节": "缺失值处理", "处理数量": 0, "说明": "保留缺失值，在具体分析时自动忽略缺失。"})

    summary = {
        "原始行数": raw_rows,
        "清洗后行数": len(clean),
        "原始字段数": len(raw_df.columns),
        "清洗后字段数": len(clean.columns),
        "重复行数": dup,
        "缺失单元格数": int(clean.isna().sum().sum())
    }
    return clean, pd.DataFrame(reports), summary, numeric_report


# ============================================================
# 2. 图表与分析函数
# ============================================================

def money_fmt(x):
    if pd.isna(x):
        return "-"
    try:
        x = float(x)
    except Exception:
        return str(x)
    if abs(x) >= 100000000:
        return f"{x / 100000000:.2f}亿"
    if abs(x) >= 10000:
        return f"{x / 10000:.2f}万"
    return f"{x:,.2f}"


def pct_fmt(x):
    if pd.isna(x):
        return "-"
    return f"{float(x):.2%}"


def kpi_card(title, value, note=""):
    st.markdown(f"""
    <div class="metric-card">
        <div class="metric-title">{title}</div>
        <div class="metric-value">{value}</div>
        <div class="metric-note">{note}</div>
    </div>
    """, unsafe_allow_html=True)


def chart_layout(fig, height=420):
    height = max(height, 500)
    fig.update_layout(
        height=height,
        template="plotly_white",
        margin=dict(l=34, r=34, t=76, b=56),
        font=dict(size=15, family="Microsoft YaHei, PingFang SC, Arial"),
        title=dict(font=dict(size=23, family="Microsoft YaHei", color="#10213F"), x=0.02, xanchor="left"),
        legend=dict(font=dict(size=14), bgcolor="rgba(255,255,255,0.72)", bordercolor="#E5EBF3", borderwidth=1),
        paper_bgcolor="#FFFFFF",
        plot_bgcolor="#FFFFFF",
        hoverlabel=dict(font_size=15, bgcolor="#FFFFFF", bordercolor="#DDE6F1")
    )
    fig.update_xaxes(
        title_font=dict(size=17, color="#53657E"),
        tickfont=dict(size=14, color="#66758C"),
        gridcolor="#E8EEF6",
        zerolinecolor="#DDE6F1",
        showline=True,
        linecolor="#DDE6F1"
    )
    fig.update_yaxes(
        title_font=dict(size=17, color="#53657E"),
        tickfont=dict(size=14, color="#66758C"),
        gridcolor="#E8EEF6",
        zerolinecolor="#DDE6F1",
        showline=True,
        linecolor="#DDE6F1"
    )
    return fig


def to_finite_numeric_series(df, col, aggressive=True):
    """把任意字段稳健转换为可绘图的有限数值序列，避免 Plotly 拿到 object/Inf/空序列后画空图。"""
    if col not in df.columns:
        return pd.Series(dtype="float64")
    s = df[col]
    if pd.api.types.is_numeric_dtype(s):
        out = pd.to_numeric(s, errors="coerce")
    else:
        out = parse_numeric_series(s, aggressive=aggressive)
    out = pd.Series(out, index=df.index).replace([np.inf, -np.inf], np.nan)
    return out.dropna().astype(float)


def to_finite_numeric_frame(df, cols, aggressive=True):
    """批量把字段转换为有限数值，用于散点、相关矩阵、异常检测等。"""
    out = pd.DataFrame(index=df.index)
    for c in cols:
        if c not in df.columns:
            continue
        if pd.api.types.is_numeric_dtype(df[c]):
            s = pd.to_numeric(df[c], errors="coerce")
        else:
            s = parse_numeric_series(df[c], aggressive=aggressive)
        s = pd.Series(s, index=df.index).replace([np.inf, -np.inf], np.nan)
        out[c] = s.astype(float)
    return out


def has_enough_numeric(df, col, min_count=1, require_variation=False):
    s = to_finite_numeric_series(df, col)
    if len(s) < min_count:
        return False
    if require_variation and s.nunique(dropna=True) <= 1:
        return False
    return True


def show_no_chart_reason(title, reason):
    """统一的非空白提示：没有可画数据时只提示原因，不再展示空坐标轴。"""
    st.warning(f"“{title}”暂无法绘制：{reason}")


def figure_has_data(fig):
    """判断 figure 是否真的有可展示的数据点，避免出现只有坐标轴的空图。"""
    try:
        for tr in fig.data:
            for attr in ["x", "y", "z", "values", "r"]:
                vals = getattr(tr, attr, None)
                if vals is None:
                    continue
                arr = np.asarray(vals, dtype=object)
                if arr.size == 0:
                    continue
                # z 可能是二维
                flat = arr.ravel()
                non_empty = [v for v in flat if v is not None and str(v) != "nan"]
                if len(non_empty) > 0:
                    return True
    except Exception:
        return True
    return False


def safe_plotly_chart(fig, title="", height=500):
    fig = chart_layout(fig, height)
    if not figure_has_data(fig):
        show_no_chart_reason(title or "图表", "当前字段没有形成有效图形数据。请更换字段或检查数值转换结果。")
        return False
    st.plotly_chart(fig, use_container_width=True)
    return True



def safe_numeric_cols(df, cols):
    """不只看 dtype，而是重新尝试数值解析，防止 Excel/CSV 数字被当成文本导致图表空白。"""
    valid = []
    for c in cols:
        if c not in df.columns or is_id_like(c):
            continue
        s = to_finite_numeric_series(df, c)
        if len(s) > 0:
            # 同步回 df 里的数值列，后续所有分析都使用干净数值
            try:
                df[c] = to_finite_numeric_frame(df, [c])[c]
            except Exception:
                pass
            valid.append(c)
    return valid


def bar_chart(df, x, y, title):
    if x not in df.columns or y not in df.columns:
        show_no_chart_reason(title, "字段不存在。")
        return
    plot = df[[x, y]].copy()
    plot[y] = pd.to_numeric(plot[y], errors="coerce").replace([np.inf, -np.inf], np.nan)
    plot = plot.dropna(subset=[x, y])
    if plot.empty:
        show_no_chart_reason(title, f"{y} 没有有效数值，或 {x} 没有有效分组。")
        return

    # 类别过多时保留前 30，避免图表挤压；如果传入前面已经 topN，这里不会改变结果
    if len(plot) > 30:
        plot = plot.sort_values(y, ascending=False).head(30)

    horizontal = len(plot) >= 8
    fig = go.Figure()
    if horizontal:
        plot = plot.sort_values(y, ascending=True)
        fig.add_trace(go.Bar(
            x=plot[y].astype(float).tolist(),
            y=plot[x].astype(str).tolist(),
            orientation="h",
            text=[money_fmt(v) for v in plot[y]],
            textposition="outside",
            marker=dict(line=dict(width=0))
        ))
        fig.update_layout(title=title, xaxis_title=y, yaxis_title=x)
    else:
        fig.add_trace(go.Bar(
            x=plot[x].astype(str).tolist(),
            y=plot[y].astype(float).tolist(),
            text=[money_fmt(v) for v in plot[y]],
            textposition="outside",
            marker=dict(line=dict(width=0))
        ))
        fig.update_layout(title=title, xaxis_title=x, yaxis_title=y)
    safe_plotly_chart(fig, title, 520)


def line_chart(df, x, y, title):
    if x not in df.columns or y not in df.columns:
        show_no_chart_reason(title, "字段不存在。")
        return
    plot = df[[x, y]].copy()
    plot[y] = pd.to_numeric(plot[y], errors="coerce").replace([np.inf, -np.inf], np.nan)
    plot = plot.dropna(subset=[x, y])
    if plot.empty:
        show_no_chart_reason(title, f"{y} 没有有效数值，或 {x} 没有有效取值。")
        return

    x_values = plot[x].astype(str).tolist()
    y_values = plot[y].astype(float).tolist()
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=x_values,
        y=y_values,
        mode="lines+markers",
        line=dict(width=4),
        marker=dict(size=8, line=dict(width=1, color="#FFFFFF")),
        hovertemplate=f"{x}=%{{x}}<br>{y}=%{{y:,.2f}}<extra></extra>"
    ))
    fig.update_layout(title=title, xaxis_title=x, yaxis_title=y, hovermode="x unified")

    # 时间点较多时只显示少量横轴刻度，避免密密麻麻挤在一起
    if len(x_values) > 12:
        step = max(1, len(x_values) // 8)
        tickvals = [x_values[i] for i in range(0, len(x_values), step)]
        if x_values[-1] not in tickvals:
            tickvals.append(x_values[-1])
        fig.update_xaxes(type="category", tickmode="array", tickvals=tickvals, tickangle=0)

    safe_plotly_chart(fig, title, 540)

def histogram_chart(df, col, title):
    s = to_finite_numeric_series(df, col)
    if s.empty:
        show_no_chart_reason(title, f"{col} 没有可解析为数值的有效数据。")
        return

    values = s.astype(float).tolist()
    bins = min(45, max(8, int(np.sqrt(len(values)) * 2)))
    fig = go.Figure()
    fig.add_trace(go.Histogram(
        x=values,
        nbinsx=bins,
        opacity=0.88,
        marker=dict(line=dict(width=0))
    ))
    fig.update_layout(title=title, xaxis_title=col, yaxis_title="记录数", bargap=0.05)
    safe_plotly_chart(fig, title, 520)


def box_chart(df, col, title):
    s = to_finite_numeric_series(df, col)
    if s.empty:
        show_no_chart_reason(title, f"{col} 没有可解析为数值的有效数据。")
        return
    fig = go.Figure()
    fig.add_trace(go.Box(
        y=s.astype(float).tolist(),
        boxpoints="outliers",
        name=col,
        marker=dict(size=5, opacity=0.55)
    ))
    fig.update_layout(title=title, yaxis_title=col)
    safe_plotly_chart(fig, title, 520)


def recommend_fit_method(df, x, y, group_col=None):
    use_cols = [x, y] + ([group_col] if group_col else [])
    plot = df[use_cols].copy()
    num = to_finite_numeric_frame(plot, [x, y])
    plot[x] = num[x] if x in num.columns else np.nan
    plot[y] = num[y] if y in num.columns else np.nan
    plot = plot.dropna(subset=[x, y])
    if len(plot) < 30:
        return "全局OLS（直线）", "样本量较少，建议先使用全局OLS观察整体方向。"
    group_n = plot[group_col].nunique() if group_col and group_col in plot.columns else 0
    corr = abs(plot[[x, y]].corr().iloc[0, 1]) if plot[x].nunique() > 1 and plot[y].nunique() > 1 else 0
    if group_n >= 2 and group_n <= 12:
        return "按维度分组OLS（分段直线）", f"当前存在 {group_n} 个有效分组，不同群体可能有不同斜率，建议使用分组OLS比较差异。"
    if corr < 0.25 and len(plot) >= 80:
        return "LOWESS（局部加权）", "整体线性相关较弱，可能存在非线性或局部拐点，建议使用LOWESS观察局部变化。"
    if len(plot) >= 80:
        return "多项式回归（二阶）", "样本量较充足，可尝试二阶多项式观察是否存在先升后降或先降后升关系。"
    return "全局OLS（直线）", "当前数据适合先用全局OLS作为整体趋势参考。"


def _add_ols_line(fig, x_values, y_values, name, dash=None):
    temp = pd.DataFrame({"x": pd.to_numeric(pd.Series(x_values), errors="coerce"),
                         "y": pd.to_numeric(pd.Series(y_values), errors="coerce")}).replace([np.inf, -np.inf], np.nan).dropna()
    if len(temp) < 3 or temp["x"].nunique() <= 1:
        return
    coef = np.polyfit(temp["x"], temp["y"], 1)
    xs = np.linspace(temp["x"].min(), temp["x"].max(), 100)
    ys = coef[0] * xs + coef[1]
    fig.add_trace(go.Scatter(x=xs.tolist(), y=ys.tolist(), mode="lines", name=name, line=dict(width=3, dash=dash or "solid")))


def _add_poly2_line(fig, x_values, y_values, name):
    temp = pd.DataFrame({"x": pd.to_numeric(pd.Series(x_values), errors="coerce"),
                         "y": pd.to_numeric(pd.Series(y_values), errors="coerce")}).replace([np.inf, -np.inf], np.nan).dropna()
    if len(temp) < 5 or temp["x"].nunique() <= 2:
        return
    coef = np.polyfit(temp["x"], temp["y"], 2)
    xs = np.linspace(temp["x"].min(), temp["x"].max(), 160)
    ys = coef[0] * xs ** 2 + coef[1] * xs + coef[2]
    fig.add_trace(go.Scatter(x=xs.tolist(), y=ys.tolist(), mode="lines", name=name, line=dict(width=3, dash="dot")))


def _add_lowess_line(fig, x_values, y_values, name):
    temp = pd.DataFrame({"x": pd.to_numeric(pd.Series(x_values), errors="coerce"),
                         "y": pd.to_numeric(pd.Series(y_values), errors="coerce")}).replace([np.inf, -np.inf], np.nan).dropna().sort_values("x")
    if len(temp) < 10 or temp["x"].nunique() <= 5:
        return
    try:
        from statsmodels.nonparametric.smoothers_lowess import lowess
        smoothed = lowess(temp["y"], temp["x"], frac=0.28, return_sorted=True)
        fig.add_trace(go.Scatter(x=smoothed[:, 0].tolist(), y=smoothed[:, 1].tolist(), mode="lines", name=name, line=dict(width=4)))
    except Exception:
        try:
            temp["_bin"] = pd.qcut(temp["x"], q=min(20, max(5, len(temp) // 20)), duplicates="drop")
            binned = temp.groupby("_bin", observed=True).agg(x=("x", "mean"), y=("y", "mean")).dropna()
            fig.add_trace(go.Scatter(x=binned["x"].tolist(), y=binned["y"].tolist(), mode="lines+markers", name=name, line=dict(width=4)))
        except Exception:
            return


def scatter_chart(df, x, y, color_col=None, fit_method="全局OLS（直线）"):
    if x not in df.columns or y not in df.columns:
        show_no_chart_reason(f"{x} 与 {y} 的关系", "字段不存在。")
        return np.nan

    cols = [x, y] + ([color_col] if color_col and color_col in df.columns else [])
    plot = df[cols].copy()
    num = to_finite_numeric_frame(plot, [x, y])
    if x not in num.columns or y not in num.columns:
        show_no_chart_reason(f"{x} 与 {y} 的关系", "横轴或纵轴无法解析为有效数值。")
        return np.nan
    plot[x] = num[x]
    plot[y] = num[y]
    plot = plot.dropna(subset=[x, y])
    if plot.empty:
        show_no_chart_reason(f"{x} 与 {y} 的关系", "两个字段没有足够的共同有效数值。")
        return np.nan

    if len(plot) > 2500:
        plot = plot.sample(2500, random_state=42)

    fig = go.Figure()
    if color_col and color_col in plot.columns:
        groups = list(plot.groupby(color_col, dropna=False))
        for idx, (name, g) in enumerate(groups):
            fig.add_trace(go.Scatter(
                x=g[x].astype(float).tolist(),
                y=g[y].astype(float).tolist(),
                mode="markers",
                name=str(name),
                marker=dict(size=6, opacity=0.48),
                showlegend=idx < 20
            ))
    else:
        fig.add_trace(go.Scatter(
            x=plot[x].astype(float).tolist(),
            y=plot[y].astype(float).tolist(),
            mode="markers",
            name="样本点",
            marker=dict(size=6, opacity=0.50)
        ))

    corr = plot[[x, y]].corr().iloc[0, 1] if len(plot) >= 3 and plot[x].nunique() > 1 and plot[y].nunique() > 1 else np.nan

    if fit_method == "全局OLS（直线）":
        _add_ols_line(fig, plot[x], plot[y], "全局OLS趋势线", dash="dash")
    elif fit_method == "按维度分组OLS（分段直线）":
        if color_col and color_col in plot.columns and plot[color_col].nunique(dropna=False) <= 15:
            for name, g in plot.groupby(color_col, dropna=False):
                _add_ols_line(fig, g[x], g[y], f"{name} OLS")
        else:
            st.info("当前未选择合适分组字段，或分组数量过多，已改用全局OLS。")
            _add_ols_line(fig, plot[x], plot[y], "全局OLS趋势线", dash="dash")
    elif fit_method == "LOWESS（局部加权）":
        _add_lowess_line(fig, plot[x], plot[y], "LOWESS局部趋势")
    elif fit_method == "多项式回归（二阶）":
        _add_poly2_line(fig, plot[x], plot[y], "二阶多项式趋势")

    fig.update_layout(title=f"{x} 与 {y} 的关系", xaxis_title=x, yaxis_title=y, hovermode="closest")
    safe_plotly_chart(fig, f"{x} 与 {y} 的关系", 540)
    return corr


def heatmap_chart(pivot, title, colorscale="Blues"):
    if pivot is None or pivot.empty:
        st.warning("当前交叉维度下没有可用于绘制热力图的数据。")
        return
    plot = pivot.copy()
    plot = plot.apply(pd.to_numeric, errors="coerce").replace([np.inf, -np.inf], np.nan)
    plot = plot.dropna(how="all", axis=0).dropna(how="all", axis=1)
    if plot.empty:
        show_no_chart_reason(title, "透视表没有有效数值。")
        return

    z = plot.fillna(0).astype(float).values.tolist()
    x_labels = [str(c) for c in plot.columns.tolist()]
    y_labels = [str(i) for i in plot.index.tolist()]

    fig = go.Figure(data=go.Heatmap(
        z=z,
        x=x_labels,
        y=y_labels,
        colorscale=colorscale,
        colorbar=dict(title="数值"),
        hovertemplate="%{y} × %{x}<br>数值=%{z:,.2f}<extra></extra>"
    ))

    # 单元格不太多时加文本，增强可读性；避免过多文字把图压乱
    if len(x_labels) * len(y_labels) <= 100:
        for i, yy in enumerate(y_labels):
            for j, xx in enumerate(x_labels):
                fig.add_annotation(
                    x=xx, y=yy, text=money_fmt(plot.iloc[i, j]),
                    showarrow=False, font=dict(size=12, color="#17324D")
                )

    fig.update_layout(
        title=title,
        xaxis_title=str(pivot.columns.name) if pivot.columns.name else "列维度",
        yaxis_title=str(pivot.index.name) if pivot.index.name else "行维度",
        xaxis=dict(type="category", tickangle=-30, automargin=True),
        yaxis=dict(type="category", automargin=True)
    )
    safe_plotly_chart(fig, title, 560)


def corr_heatmap_chart(df, numeric_cols):
    # 优先用用户选中的数值字段；如果不足，自动扫描全表中能解析为数值的字段
    scan_cols = list(dict.fromkeys(list(numeric_cols) + list(df.columns)))
    numeric_df = to_finite_numeric_frame(df, [c for c in scan_cols if c in df.columns and not is_id_like(c)])
    cols = []
    for c in numeric_df.columns:
        valid = numeric_df[c].dropna()
        if len(valid) >= 3 and valid.nunique(dropna=True) > 1:
            cols.append(c)
    if len(cols) < 2:
        st.warning("可用于相关性矩阵的数值字段不足。请至少选择两个有效数值字段，且字段不能全为同一个值。")
        return

    corr = numeric_df[cols].corr(method="pearson").replace([np.inf, -np.inf], np.nan)
    corr = corr.dropna(how="all", axis=0).dropna(how="all", axis=1)
    if corr.empty or corr.shape[0] < 2 or corr.shape[1] < 2:
        show_no_chart_reason("数值指标相关性矩阵", "相关系数无法计算，可能是字段有效值过少或没有波动。")
        return

    for c in corr.columns:
        if c in corr.index:
            corr.loc[c, c] = 1.0

    z = corr.fillna(0).astype(float).values.tolist()
    x_labels = corr.columns.astype(str).tolist()
    y_labels = corr.index.astype(str).tolist()

    fig = go.Figure(data=go.Heatmap(
        z=z,
        x=x_labels,
        y=y_labels,
        zmin=-1,
        zmax=1,
        colorscale="RdBu",
        reversescale=True,
        colorbar=dict(title="相关系数"),
        hovertemplate="%{y} 与 %{x}<br>相关系数=%{z:.3f}<extra></extra>"
    ))

    if len(x_labels) * len(y_labels) <= 144:
        for i, yy in enumerate(y_labels):
            for j, xx in enumerate(x_labels):
                fig.add_annotation(
                    x=xx, y=yy, text=f"{corr.iloc[i, j]:.2f}",
                    showarrow=False, font=dict(size=12, color="#17324D")
                )

    fig.update_layout(
        title="数值指标相关性矩阵",
        xaxis_title="数值字段",
        yaxis_title="数值字段",
        xaxis=dict(type="category", tickangle=-30, automargin=True),
        yaxis=dict(type="category", automargin=True, autorange="reversed")
    )
    safe_plotly_chart(fig, "数值指标相关性矩阵", 600)

    pairs = []
    for i, a in enumerate(corr.columns):
        for j, b in enumerate(corr.index):
            if j <= i:
                continue
            val = corr.loc[b, a]
            if pd.notna(val) and abs(val) >= 0.6:
                pairs.append((a, b, val))
    if pairs:
        pairs = sorted(pairs, key=lambda t: abs(t[2]), reverse=True)[:5]
        msg = "；".join([f"{a} 与 {b} 的相关系数为 {v:.2f}" for a, b, v in pairs])
        st.info(f"系统识别到较强相关关系：{msg}。这些字段可能是影响主指标变化的重要线索。")
    else:
        st.info("当前数值字段之间未发现特别强的线性相关关系。接近 1 表示强正相关，接近 -1 表示强负相关，接近 0 表示线性关系较弱。")

def periodize(df, date_col):
    temp = df.copy()
    temp["_period"] = pd.to_datetime(temp[date_col], errors="coerce").dt.to_period("M").astype(str)
    temp = temp[temp["_period"].notna()]
    return temp


def build_trend(df, date_col, metric):
    temp = periodize(df, date_col)
    if temp.empty:
        return pd.DataFrame(columns=["期间", metric])
    trend = temp.groupby("_period", dropna=True)[metric].sum().reset_index()
    trend = trend.rename(columns={"_period": "期间"}).sort_values("期间")
    return trend



def trend_interpretation(trend, metric):
    if trend is None or len(trend) < 2:
        return "当前数据不足以判断趋势。建议至少提供两个以上时间周期的数据。"
    trend = trend.dropna(subset=[metric]).copy()
    if len(trend) < 2:
        return "当前主指标在时间维度上的有效数值不足，暂无法形成稳定趋势判断。"
    first, last = trend[metric].iloc[0], trend[metric].iloc[-1]
    change = (last-first)/abs(first) if first != 0 else np.nan
    diffs = trend[metric].diff().dropna()
    up_count, down_count, flat_count = int((diffs>0).sum()), int((diffs<0).sum()), int((diffs==0).sum())
    recent = trend.tail(min(3, len(trend)))
    recent_change = np.nan
    if len(recent) >= 2 and recent[metric].iloc[0] != 0:
        recent_change = (recent[metric].iloc[-1] - recent[metric].iloc[0]) / abs(recent[metric].iloc[0])
    max_row = trend.loc[trend[metric].idxmax()]
    min_row = trend.loc[trend[metric].idxmin()]
    mean_val = trend[metric].mean()
    cv = trend[metric].std()/abs(mean_val) if mean_val != 0 else np.nan
    if pd.isna(change): main_desc = "整体变化方向暂无法计算"
    elif change > 0.1: main_desc = f"整体呈上升趋势，末期较初期增长约 {change:.2%}"
    elif change < -0.1: main_desc = f"整体呈下降趋势，末期较初期下降约 {abs(change):.2%}"
    else: main_desc = "整体较为平稳，期初与期末差异不大"
    if not pd.isna(recent_change):
        if recent_change > 0.08: recent_desc = f"最近 {len(recent)} 期继续上行，累计增长约 {recent_change:.2%}"
        elif recent_change < -0.08: recent_desc = f"最近 {len(recent)} 期出现回落，累计下降约 {abs(recent_change):.2%}"
        else: recent_desc = f"最近 {len(recent)} 期变化幅度较小，短期表现相对稳定"
    else: recent_desc = "短期变化暂无法计算"
    if pd.isna(cv): vol_desc = "波动性暂无法计算"
    elif cv < 0.15: vol_desc = "整体波动较低，指标运行较稳定"
    elif cv < 0.45: vol_desc = "存在一定波动，需要结合高低点进一步观察"
    else: vol_desc = "波动较强，建议重点关注异常月份或阶段性变化"
    detail = f"在全部相邻周期中，上升 {up_count} 次、下降 {down_count} 次、持平 {flat_count} 次。"
    high_low = f"最高值出现在 {max_row['期间']}，为 {money_fmt(max_row[metric])}；最低值出现在 {min_row['期间']}，为 {money_fmt(min_row[metric])}。"
    return f"{main_desc}。{recent_desc}。{detail}{high_low}{vol_desc}。"

def dimension_summary(df, dim, metric):
    # 主指标如果是折扣/比率/单价等水平型指标，核心排序口径用均值；
    # 如果是销售额/数量/利润等规模型指标，核心排序口径用合计。
    method = aggregation_method_for_metric(metric)
    g = df.groupby(dim, dropna=False)[metric].agg(["sum", "mean", "count"]).reset_index()
    g.columns = [dim, f"{metric}合计", f"{metric}均值", "记录数"]
    sort_col = f"{metric}均值" if method == "mean" else f"{metric}合计"
    return g.sort_values(sort_col, ascending=False)


def dimension_importance(df, dimensions, metric):
    rows = []
    total_var = df[metric].var()
    for dim in dimensions:
        if dim not in df.columns:
            continue
        g = df.groupby(dim, dropna=False)[metric].mean()
        score = 0 if pd.isna(total_var) or total_var == 0 else min(float(g.var() / total_var), 1)
        rows.append({"维度": dim, "差异解释度": score, "维度取值数": int(df[dim].nunique(dropna=True))})
    return pd.DataFrame(rows).sort_values("差异解释度", ascending=False) if rows else pd.DataFrame()



def contribution_analysis(df, date_col, dim, metric, start_period=None, end_period=None):
    if not date_col or not dim: return pd.DataFrame()
    temp = periodize(df, date_col)
    periods = sorted(temp['_period'].dropna().unique().tolist())
    if len(periods) < 2: return pd.DataFrame()
    if start_period is None or end_period is None:
        p0, p1 = periods[-2], periods[-1]
    else:
        p0, p1 = start_period, end_period
    if p0 == p1: return pd.DataFrame()
    a = temp[temp['_period']==p0].groupby(dim)[metric].sum()
    b = temp[temp['_period']==p1].groupby(dim)[metric].sum()
    comp = pd.concat([a,b], axis=1).fillna(0)
    comp.columns = [p0,p1]
    comp['变化额'] = comp[p1]-comp[p0]
    total_change = comp['变化额'].sum()
    comp['贡献占比'] = comp['变化额']/total_change if total_change != 0 else np.nan
    return comp.reset_index().sort_values('变化额', ascending=False)

def relationship_explanation(x, y, corr):
    if pd.isna(corr):
        return "当前两个指标缺失较多，暂无法计算稳定的相关性。"
    strength = "较强" if abs(corr) >= 0.6 else ("中等" if abs(corr) >= 0.3 else "较弱")
    direction = "正向" if corr > 0 else ("反向" if corr < 0 else "无明显")
    return f"相关系数为 {corr:.3f}，说明 {x} 与 {y} 的线性相关性{strength}，方向为{direction}关系。散点图主要用于观察两个数值指标之间是否存在同向、反向或异常偏离关系。"




def correlation_summary_tables(df, numeric_cols):
    """汇总所有经营数值指标之间的相关系数，便于用户直接查看整体关系。"""
    scan_cols = [c for c in numeric_cols if c in df.columns and not is_id_like(c)]
    numeric_df = to_finite_numeric_frame(df, scan_cols)
    valid_cols = []
    for c in numeric_df.columns:
        s = numeric_df[c].dropna()
        if len(s) >= 3 and s.nunique(dropna=True) > 1:
            valid_cols.append(c)
    if len(valid_cols) < 2:
        return pd.DataFrame(), pd.DataFrame()

    corr = numeric_df[valid_cols].corr(method="pearson").replace([np.inf, -np.inf], np.nan)
    corr = corr.dropna(how="all", axis=0).dropna(how="all", axis=1)

    rows = []
    cols = corr.columns.tolist()
    for i in range(len(cols)):
        for j in range(i + 1, len(cols)):
            a, b = cols[i], cols[j]
            val = corr.loc[a, b]
            if pd.isna(val):
                continue
            abs_val = abs(float(val))
            if abs_val >= 0.7:
                strength = "强相关"
            elif abs_val >= 0.4:
                strength = "中等相关"
            elif abs_val >= 0.2:
                strength = "弱相关"
            else:
                strength = "相关较弱"
            direction = "正相关" if val > 0 else ("负相关" if val < 0 else "无明显方向")
            rows.append({
                "指标A": a,
                "指标B": b,
                "相关系数": round(float(val), 3),
                "绝对值": round(abs_val, 3),
                "方向": direction,
                "强度判断": strength,
                "经营含义提示": f"{a} 与 {b} 呈{strength}、{direction}，可作为后续驱动分析或交叉核查线索。"
            })
    pair_df = pd.DataFrame(rows).sort_values("绝对值", ascending=False) if rows else pd.DataFrame()
    return corr.round(3), pair_df


def parse_markdown_tables(text):
    """将大模型输出中的 Markdown 表格解析成 DataFrame，避免页面显示 |---|---|。"""
    if not text:
        return [], ""
    lines = str(text).splitlines()
    blocks, normal_lines = [], []
    i = 0
    sep_pat = r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$"
    while i < len(lines):
        line = lines[i].strip()
        if "|" in line and i + 1 < len(lines) and re.match(sep_pat, lines[i + 1].strip()):
            table_lines = [lines[i].strip(), lines[i + 1].strip()]
            i += 2
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            header = [x.strip() for x in table_lines[0].strip("|").split("|")]
            data = []
            for row_line in table_lines[2:]:
                cells = [x.strip() for x in row_line.strip("|").split("|")]
                if len(cells) < len(header):
                    cells += [""] * (len(header) - len(cells))
                data.append(cells[:len(header)])
            df_table = pd.DataFrame(data, columns=header)
            before = "\n".join(normal_lines).strip()
            blocks.append((before, df_table))
            normal_lines = []
        else:
            normal_lines.append(lines[i])
            i += 1
    return blocks, "\n".join(normal_lines).strip()


def render_ai_answer_pretty(ans):
    """美化大模型解释结果：正文正常显示，Markdown表格转成真正表格。"""
    blocks, tail = parse_markdown_tables(ans)
    if not blocks:
        st.markdown(ans)
        return
    for before, table_df in blocks:
        if before:
            st.markdown(before)
        st.markdown('<div class="ai-table-card"><div class="ai-section-title">结构化核查清单</div></div>', unsafe_allow_html=True)
        st.dataframe(table_df, use_container_width=True, hide_index=True)
    if tail:
        st.markdown(tail)


def _sparse_period_ticks(periods, max_ticks=8):
    periods = [str(x) for x in periods]
    if len(periods) <= max_ticks:
        return periods
    step = max(1, int(np.ceil(len(periods) / max_ticks)))
    vals = [periods[i] for i in range(0, len(periods), step)]
    if periods[-1] not in vals:
        vals.append(periods[-1])
    return vals


def animated_trend_chart(trend, metric, title=None):
    """
    更清晰的动态趋势演示：
    1. 每一帧展示从起点到当前周期的累计轨迹；
    2. 横轴只显示少量区间刻度，避免时间标签密集重叠；
    3. 单独高亮当前点，增强趋势演变可读性。
    """
    if trend is None or len(trend) < 3 or "期间" not in trend.columns or metric not in trend.columns:
        st.info("当前未生成动态图：需要有效日期字段，并且时间周期不少于 3 个。")
        return

    plot = trend[["期间", metric]].dropna().copy()
    plot[metric] = pd.to_numeric(plot[metric], errors="coerce")
    plot = plot.dropna()
    if len(plot) < 3:
        st.info("当前未生成动态图：当前主指标在时间维度上的有效点不足。")
        return

    periods = plot["期间"].astype(str).tolist()
    values = plot[metric].astype(float).tolist()
    ymin, ymax = min(values), max(values)
    pad = (ymax - ymin) * 0.15 if ymax != ymin else max(abs(ymax) * 0.1, 1)
    tickvals = _sparse_period_ticks(periods, max_ticks=8)

    frames = []
    for i in range(1, len(plot) + 1):
        frames.append(go.Frame(
            data=[
                go.Scatter(
                    x=periods[:i],
                    y=values[:i],
                    mode="lines",
                    fill="tozeroy",
                    line=dict(width=5),
                    name=f"{metric}累计轨迹",
                    hovertemplate="期间=%{x}<br>数值=%{y:,.2f}<extra></extra>"
                ),
                go.Scatter(
                    x=[periods[i - 1]],
                    y=[values[i - 1]],
                    mode="markers+text",
                    marker=dict(size=16, symbol="circle"),
                    text=[money_fmt(values[i - 1])],
                    textposition="top center",
                    name="当前周期",
                    hovertemplate="当前周期=%{x}<br>数值=%{y:,.2f}<extra></extra>"
                )
            ],
            name=periods[i - 1]
        ))

    fig = go.Figure(
        data=[
            go.Scatter(
                x=[periods[0]],
                y=[values[0]],
                mode="lines",
                fill="tozeroy",
                line=dict(width=5),
                name=f"{metric}累计轨迹"
            ),
            go.Scatter(
                x=[periods[0]],
                y=[values[0]],
                mode="markers+text",
                marker=dict(size=16),
                text=[money_fmt(values[0])],
                textposition="top center",
                name="当前周期"
            )
        ],
        frames=frames
    )

    slider_steps = []
    label_step = max(1, int(np.ceil(len(periods) / 8)))
    for i, p in enumerate(periods):
        slider_steps.append({
            "args": [[p], {"frame": {"duration": 0, "redraw": True}, "mode": "immediate"}],
            "label": p if (i % label_step == 0 or i == len(periods) - 1) else "",
            "method": "animate"
        })

    fig.update_layout(
        title=title or f"{metric}动态趋势演示",
        xaxis=dict(
            title="期间",
            type="category",
            categoryorder="array",
            categoryarray=periods,
            tickmode="array",
            tickvals=tickvals,
            tickangle=0
        ),
        yaxis=dict(title=metric, range=[ymin - pad, ymax + pad]),
        hovermode="x unified",
        updatemenus=[{
            "type": "buttons",
            "showactive": False,
            "x": 0.02,
            "y": 1.15,
            "buttons": [
                {"label": "播放", "method": "animate", "args": [None, {"frame": {"duration": 520, "redraw": True}, "fromcurrent": True}]},
                {"label": "暂停", "method": "animate", "args": [[None], {"frame": {"duration": 0, "redraw": False}, "mode": "immediate"}]}
            ]
        }],
        sliders=[{
            "steps": slider_steps,
            "currentvalue": {"prefix": "当前周期：", "font": {"size": 15}},
            "x": 0.08,
            "len": 0.86,
            "pad": {"t": 45, "b": 10}
        }]
    )
    safe_plotly_chart(fig, title or f"{metric}动态趋势演示", 590)


def clean_md_inline(text):
    """清理 AI 输出和 Word 文本中的 Markdown 标记、异常空格和多余引号。"""
    text = str(text or "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text.strip())
    text = text.replace("---", "")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+([，。；：！？、）】》])", r"\1", text)
    text = re.sub(r"([（【《])\s+", r"\1", text)
    text = text.replace("“ ", "“").replace(" ”", "”").replace("‘ ", "‘").replace(" ’", "’")
    text = text.replace("\u200b", "").strip()
    return text


def clean_report_sentence(text):
    """用于写入 Word 的最终文本清洗，避免出现 Markdown 残留和引号断裂。"""
    text = clean_md_inline(text)
    text = re.sub(r"^[\-•]\s*", "", text)
    text = text.replace("|", "，")
    text = re.sub(r"\s*，\s*", "，", text)
    text = re.sub(r"\s*。\s*", "。", text)
    return text.strip()


def clean_focus_list_for_report(focus_list, date_col):
    """无有效日期字段时，避免报告仍显示“趋势变化”关注重点。"""
    cleaned = []
    for item in (focus_list or []):
        s = str(item).strip()
        if not date_col and any(k in s for k in ["趋势", "时间", "周期"]):
            continue
        if s and s not in cleaned:
            cleaned.append(s)
    if not date_col and "结构分布" not in cleaned:
        cleaned.append("结构分布")
    return cleaned


def is_report_meta_line(text):
    """判断 AI 输出中不适合放入第五部分的报告标题/元信息。"""
    s = clean_report_sentence(text)
    if not s:
        return True
    meta_patterns = [
        r"^(经营分析补充解读|经营分析决策简报|智策经营|AI增强版|管理层经营简报)",
        r"^数据周期[:：]", r"^主指标[:：]", r"^分析对象[:：]", r"^报告生成",
        r"^本报告基于", r"^本次简报关注重点", r"^一、", r"^二、", r"^三、", r"^四、", r"^五、",
        r"^\d+\.\d+", r"^第[一二三四五六七八九十]+部分"
    ]
    return any(re.search(p, s) for p in meta_patterns)
def add_dataframe_to_docx(doc, df_table, max_rows=12):
    if df_table is None or len(df_table) == 0:
        return
    show = df_table.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, col in enumerate(show.columns):
        cell = table.rows[0].cells[j]
        run = cell.paragraphs[0].add_run(clean_report_sentence(col))
        set_run_font(run, "黑体", 9.5, True)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(show.columns):
            value = clean_report_sentence(row[col])
            if len(value) > 95:
                value = value[:95] + "……"
            run = cells[j].paragraphs[0].add_run(value)
            set_run_font(run, "宋体", 9, False)
def add_native_bar_visual(doc, title, rows, label_col="项目", value_col="数值", max_rows=10):
    """用 Word 原生表格做条形图，避免云端导出图片时中文乱码。"""
    if rows is None or len(rows) == 0:
        return
    add_heading(doc, title, 2)
    data = pd.DataFrame(rows).copy().head(max_rows)
    data[value_col] = pd.to_numeric(data[value_col], errors="coerce").fillna(0)
    max_val = data[value_col].abs().max()
    max_val = max(max_val, 1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [label_col, value_col, "可视化条"]
    for j, h in enumerate(headers):
        run = table.rows[0].cells[j].paragraphs[0].add_run(h)
        set_run_font(run, "黑体", 10, True)

    for _, r in data.iterrows():
        cells = table.add_row().cells
        label = str(r[label_col])
        val = float(r[value_col])
        bar_len = int(abs(val) / max_val * 24)
        bar = "█" * max(1, bar_len)
        vals = [label, money_fmt(val), bar]
        for j, v in enumerate(vals):
            run = cells[j].paragraphs[0].add_run(str(v))
            set_run_font(run, "宋体", 9, False)


def add_trend_visual_table(doc, trend, metric):
    if trend is None or len(trend) == 0:
        return
    show = trend.tail(10).copy()
    show[metric] = pd.to_numeric(show[metric], errors="coerce").fillna(0)
    rows = [{"期间": r["期间"], "数值": r[metric]} for _, r in show.iterrows()]
    add_native_bar_visual(doc, f"{metric}近10期趋势可视化", rows, "期间", "数值", 10)


def add_dimension_visual_table(doc, df, dim, metric):
    if not dim or dim not in df.columns:
        return
    g = dimension_summary(df, dim, metric).head(10)
    rows = [{dim: r[dim], f"{metric}合计": r[f"{metric}合计"]} for _, r in g.iterrows()]
    add_native_bar_visual(doc, f"按{dim}的{metric}Top10可视化", rows, dim, f"{metric}合计", 10)


def add_risk_visual_table(doc, anomaly_df):
    if anomaly_df is None or len(anomaly_df) == 0 or "风险等级" not in anomaly_df.columns:
        return
    rc = anomaly_df["风险等级"].value_counts().reset_index()
    rc.columns = ["风险等级", "数量"]
    rows = [{"风险等级": r["风险等级"], "数量": r["数量"]} for _, r in rc.iterrows()]
    add_native_bar_visual(doc, "风险等级分布可视化", rows, "风险等级", "数量", 10)


def add_markdown_body_to_docx(doc, text, df, main_metric, dimensions, date_col, anomaly_df):
    """
    将 AI Markdown 正文整理成正式 Word 结构：
    - 去掉 ###、**、--- 等 Markdown 符号；
    - 一级/二级标题转为 Word 标题；
    - Markdown 表格转为 Word 表格；
    - 在对应章节中穿插原生可视化表，不再把图表集中堆到最后。
    """
    lines = str(text).splitlines()
    current_section = ""
    i = 0
    sep_pat = r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$"
    inserted_trend = False
    inserted_dim = False
    inserted_risk = False

    while i < len(lines):
        raw = lines[i].strip()
        if not raw or raw == "---":
            i += 1
            continue

        # Markdown table
        if "|" in raw and i + 1 < len(lines) and re.match(sep_pat, lines[i + 1].strip()):
            table_lines = [raw, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            header = [clean_md_inline(x.strip()) for x in table_lines[0].strip("|").split("|")]
            data = []
            for row_line in table_lines[2:]:
                cells = [clean_md_inline(x.strip()) for x in row_line.strip("|").split("|")]
                if len(cells) < len(header):
                    cells += [""] * (len(header) - len(cells))
                data.append(cells[:len(header)])
            add_dataframe_to_docx(doc, pd.DataFrame(data, columns=header), max_rows=15)
            continue

        # Headings
        heading_match = re.match(r"^#{1,6}\s*(.+)$", raw)
        cn_heading_match = re.match(r"^([一二三四五六七八九十]+、.+)$", raw)
        num_heading_match = re.match(r"^(\d+[\.、].+)$", raw)

        if heading_match or cn_heading_match or num_heading_match:
            title = heading_match.group(1) if heading_match else (cn_heading_match.group(1) if cn_heading_match else num_heading_match.group(1))
            title = clean_md_inline(title)
            if "经营分析补充解读" in title or "基于" in title and "主指标" in title:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title); set_run_font(run, "黑体", 14, True)
            else:
                current_section = title
                add_heading(doc, title, 1 if re.match(r"^[一二三四五六七八九十]+、", title) else 2)

                if ("核心经营表现" in title or "时间趋势" in title) and not inserted_trend:
                    if date_col:
                        trend = build_trend(df, date_col, main_metric)
                        if len(trend) >= 2:
                            add_trend_visual_table(doc, trend, main_metric)
                            inserted_trend = True
                    if dimensions and not inserted_dim:
                        add_dimension_visual_table(doc, df, dimensions[0], main_metric)
                        inserted_dim = True

                if ("风险识别" in title or "异常" in title) and not inserted_risk:
                    add_risk_visual_table(doc, anomaly_df)
                    inserted_risk = True
            i += 1
            continue

        # Bullets / normal paragraphs
        clean = clean_md_inline(raw)
        clean = re.sub(r"^[\-\*]\s+", "", clean)
        clean = clean.replace("✅", "【行动】").strip()
        if clean:
            add_para(doc, clean, indent=not clean.startswith(("【行动】", "注：", ">")))
        i += 1

    # 如果 AI 文本没有触发对应标题，也补充可视化，保证报告有图表类内容
    if date_col and not inserted_trend:
        trend = build_trend(df, date_col, main_metric)
        if len(trend) >= 2:
            add_trend_visual_table(doc, trend, main_metric)
    if dimensions and not inserted_dim:
        add_dimension_visual_table(doc, df, dimensions[0], main_metric)
    if anomaly_df is not None and len(anomaly_df) and not inserted_risk:
        add_risk_visual_table(doc, anomaly_df)




def extract_clean_ai_blocks(ai_text):
    """提取 AI 正文中的有效段落和表格，去掉报告题名、数据周期、一级编号标题和 Markdown 标记。"""
    text = str(ai_text or "")
    table_blocks, _ = parse_markdown_tables(text)
    lines = text.splitlines()
    paras = []
    for line in lines:
        raw = line.strip()
        if not raw or raw == "---" or "|" in raw:
            continue
        raw = re.sub(r"^#{1,6}\s*", "", raw).strip()
        raw = re.sub(r"^[\-•]\s*", "", raw).strip()
        raw = clean_report_sentence(raw)
        raw = re.sub(r"^\d+[\.、]\s*", "", raw).strip()
        if is_report_meta_line(raw):
            continue
        if len(raw) >= 12:
            paras.append(raw)
    dedup = []
    for p in paras:
        if p not in dedup:
            dedup.append(p)
    return dedup[:10], table_blocks
def add_default_action_table(doc, main_metric, dimensions):
    rows = [
        {"核查方向": "主指标异常单元", "具体内容": f"复核{main_metric}处于高位或低位的经营单元", "输出要求": "标注异常来源、责任维度和复核结论"},
        {"核查方向": "原始明细数据", "具体内容": "核查订单、项目、客户、产品等原始明细记录", "输出要求": "确认是否存在重复、缺失、口径不一致或录入错误"},
        {"核查方向": "关键维度下钻", "具体内容": f"围绕{dimensions[0] if dimensions else '主要维度'}进行分组对比", "输出要求": "识别贡献最高、波动最大或风险最集中的维度项"},
        {"核查方向": "管理动作", "具体内容": "结合风险等级设置复核优先级和跟踪周期", "输出要求": "形成可执行的整改措施和后续监控指标"},
    ]
    add_dataframe_to_docx(doc, pd.DataFrame(rows), max_rows=8)



def _report_metric_profile(df, metric):
    s = pd.to_numeric(df[metric], errors="coerce").replace([np.inf, -np.inf], np.nan).dropna()
    if len(s) == 0:
        return {
            "total": 0, "mean": 0, "median": 0, "max": 0, "min": 0,
            "cv": np.nan, "max_mean_ratio": np.nan, "skew_desc": "当前主指标有效数值不足"
        }
    total = float(s.sum())
    mean = float(s.mean())
    median = float(s.median())
    max_v = float(s.max())
    min_v = float(s.min())
    cv = float(s.std() / abs(mean)) if mean != 0 and len(s) > 1 else np.nan
    max_mean_ratio = float(max_v / mean) if mean != 0 else np.nan
    if not pd.isna(max_mean_ratio) and max_mean_ratio >= 8:
        skew_desc = "主指标呈明显长尾特征，少数高值经营单元对整体结果影响较大"
    elif not pd.isna(max_mean_ratio) and max_mean_ratio >= 4:
        skew_desc = "主指标存在一定高值集中现象，需要关注高值经营单元的来源和质量"
    else:
        skew_desc = "主指标分布相对平缓，极端高值对整体结果的影响相对有限"
    return {
        "total": total, "mean": mean, "median": median, "max": max_v, "min": min_v,
        "cv": cv, "max_mean_ratio": max_mean_ratio, "skew_desc": skew_desc
    }


def _top_dimension_insight(df, dimensions, metric):
    if not dimensions:
        return None
    dim = dimensions[0]
    if dim not in df.columns:
        return None
    g = dimension_summary(df, dim, metric)
    if len(g) == 0:
        return None
    total = float(pd.to_numeric(df[metric], errors="coerce").sum())
    top = g.iloc[0]
    top_value = float(top[f"{metric}合计"])
    share = top_value / total if total else np.nan
    return {
        "dim": dim,
        "top_name": top[dim],
        "top_value": top_value,
        "share": share,
        "table": g
    }


def _risk_profile(anomaly_df):
    if anomaly_df is None or len(anomaly_df) == 0 or "是否经营异常" not in anomaly_df.columns:
        return {"abnormal_n": 0, "unit_n": 0, "high_n": 0, "rate": 0, "high_rate": 0}
    unit_n = len(anomaly_df)
    abnormal = anomaly_df[anomaly_df["是否经营异常"]].copy()
    abnormal_n = len(abnormal)
    high_n = int((abnormal.get("风险等级", pd.Series(dtype=str)) == "高风险").sum()) if abnormal_n else 0
    return {
        "abnormal_n": abnormal_n,
        "unit_n": unit_n,
        "high_n": high_n,
        "rate": abnormal_n / unit_n if unit_n else 0,
        "high_rate": high_n / abnormal_n if abnormal_n else 0
    }


def _split_evidence(ev):
    if pd.isna(ev):
        return []
    parts = re.split(r"[;；。]\s*", str(ev))
    return [p.strip() for p in parts if p and p.strip()]


def _major_issue_from_evidence(ev, main_metric):
    parts = _split_evidence(ev)
    tags = []
    joined = "；".join(parts)
    if "负" in joined and ("利润" in joined or "收益" in joined):
        tags.append("存在负利润或盈利质量风险")
    if "折扣" in joined or "比例" in joined:
        tags.append("折扣/比例口径需复核")
    if "多项指标" in joined or "极端区间" in joined:
        tags.append("多指标同步偏离")
    if main_metric and main_metric in joined:
        tags.append(f"{main_metric}处于极端区间")
    if not tags:
        tags = parts[:2] if parts else ["异常依据需结合明细复核"]
    return "；".join(dict.fromkeys(tags))


def _action_from_evidence(ev, main_metric):
    joined = str(ev)
    actions = []
    if "折扣" in joined or "比例" in joined:
        actions.append("核查折扣政策、审批权限和录入口径")
    if "利润" in joined or "收益" in joined or "负" in joined:
        actions.append("核查成本归集、售价政策和利润明细")
    if "数量" in joined:
        actions.append("核查订单数量、拆单合单和重复记录")
    if main_metric and main_metric in joined:
        actions.append("核查高值业务单据、客户/产品来源和交易真实性")
    if not actions:
        actions.append("核查原始明细、业务口径和相关审批记录")
    return "；".join(dict.fromkeys(actions[:3]))


def _build_risk_review_table(anomaly_df, main_metric, dimensions, max_rows=5):
    """生成面向管理层的异常核查清单：说明查谁、查什么、为什么查。"""
    if anomaly_df is None or len(anomaly_df) == 0 or "是否经营异常" not in anomaly_df.columns:
        return pd.DataFrame()

    abnormal = anomaly_df[anomaly_df["是否经营异常"]].copy()
    if len(abnormal) == 0:
        return pd.DataFrame()
    if "风险得分" in abnormal.columns:
        abnormal = abnormal.sort_values("风险得分", ascending=False)
    top = abnormal.head(max_rows).copy()

    risk_cols = {
        "是否AI异常", "异常得分", "是否经营异常", "异常依据", "主指标偏离", "多指标组合偏离",
        "业务规则风险", "模型异常贡献", "风险得分", "风险等级", "记录数"
    }

    obj_cols = []
    if "_period" in top.columns:
        obj_cols.append("_period")
    for d in dimensions or []:
        if d in top.columns and d not in obj_cols:
            obj_cols.append(d)
    if not obj_cols:
        for c in top.columns:
            if c not in risk_cols and not pd.api.types.is_numeric_dtype(top[c]):
                obj_cols.append(c)
            if len(obj_cols) >= 2:
                break

    value_cols = []
    for c in [main_metric, "销售额", "数量", "折扣", "利润", "成本", "费用"]:
        if c in top.columns and c not in value_cols:
            value_cols.append(c)
    # 自动补充少量金额/数量/利润类字段
    for c in top.columns:
        if c in value_cols or c in risk_cols or c in obj_cols:
            continue
        if any(k in str(c) for k in ["销售", "收入", "金额", "利润", "成本", "费用", "数量", "折扣"]):
            if c in top.columns and pd.api.types.is_numeric_dtype(top[c]):
                value_cols.append(c)
        if len(value_cols) >= 5:
            break

    rows = []
    for i, (_, row) in enumerate(top.iterrows(), 1):
        obj_parts = []
        for c in obj_cols[:3]:
            v = row.get(c, "")
            label = "期间" if c == "_period" else c
            obj_parts.append(f"{label}={v}")
        obj = "；".join(obj_parts) if obj_parts else f"异常单元{i}"

        metric_parts = []
        for c in value_cols[:5]:
            v = row.get(c, np.nan)
            if pd.notna(v):
                metric_parts.append(f"{c}={money_fmt(v)}")
        ev = row.get("异常依据", "")
        rows.append({
            "核查对象": obj,
            "关键指标": "；".join(metric_parts) if metric_parts else "-",
            "风险等级": row.get("风险等级", "-"),
            "风险得分": row.get("风险得分", "-"),
            "主要异常问题": _major_issue_from_evidence(ev, main_metric),
            "建议核查方向": _action_from_evidence(ev, main_metric)
        })
    return pd.DataFrame(rows)


def _build_management_action_table(main_metric, dimensions, anomaly_df=None):
    dim_text = dimensions[0] if dimensions else "主要业务维度"
    rows = [
        {
            "管理关注点": "高值经营单元质量",
            "需要解决的问题": f"{main_metric}高值是否来自真实业务增长，还是由少数异常记录、重复记录或口径差异拉高",
            "建议动作": "抽取Top高值记录，核对订单/合同/客户/产品明细，确认收入确认与业务口径一致"
        },
        {
            "管理关注点": "结构集中风险",
            "需要解决的问题": f"{dim_text}下是否存在销售贡献过度集中，导致对单一渠道、客户或业务类型依赖过高",
            "建议动作": f"围绕{dim_text}进行下钻，比较记录数、金额合计、利润表现和平均折扣，区分正常规模优势与异常集中"
        },
        {
            "管理关注点": "异常风险复核",
            "需要解决的问题": "高风险对象是否同时存在高值、高折扣、负利润或多指标同步偏离",
            "建议动作": "优先核查风险得分最高的经营单元，调取原始单据、审批记录、折扣政策和成本核算表"
        }
    ]
    return pd.DataFrame(rows)


def _add_key_conclusion_box(doc, df, main_metric, dimensions, date_col, anomaly_df):
    profile = _report_metric_profile(df, main_metric)
    dim_info = _top_dimension_insight(df, dimensions, main_metric)
    risk = _risk_profile(anomaly_df)

    add_heading(doc, "一、核心结论摘要", 1)
    add_para(doc, f"本简报面向经营管理者，重点用于快速判断当前{main_metric}表现、识别主要贡献来源、定位异常风险对象，并形成后续核查与管理行动清单。", indent=True)

    conclusions = []
    conclusions.append(f"经营规模方面，{main_metric}合计为{money_fmt(profile['total'])}，均值为{money_fmt(profile['mean'])}，最大值为{money_fmt(profile['max'])}；{profile['skew_desc']}。")
    if dim_info:
        share_text = pct_fmt(dim_info["share"]) if not pd.isna(dim_info["share"]) else "-"
        conclusions.append(f"结构贡献方面，{dim_info['dim']}维度下“{dim_info['top_name']}”贡献最高，合计{money_fmt(dim_info['top_value'])}，占主指标总量约{share_text}，建议作为重点下钻对象。")
    if risk["unit_n"]:
        conclusions.append(f"风险识别方面，系统识别出{risk['abnormal_n']}个异常经营单元，其中高风险单元{risk['high_n']}个；异常结果主要用于提示优先核查对象，而不是直接替代业务判断。")
    if not date_col:
        conclusions.append("趋势条件方面，当前未选择有效日期字段，因此不形成时间趋势判断，报告转为结构分布、多维对比和异常核查分析。")

    for i, c in enumerate(conclusions, 1):
        add_para(doc, f"{i}. {c}", indent=False)


def _add_distribution_interpretation(doc, df, metric):
    profile = _report_metric_profile(df, metric)
    add_para(doc, f"图表解读：{profile['skew_desc']}。若图中大量样本集中在低值区间，同时存在少量高值记录，说明后续应重点核查高值经营单元的业务来源、客户结构和盈利质量。", indent=True)


def _add_dimension_interpretation(doc, df, dimensions, metric):
    dim_info = _top_dimension_insight(df, dimensions, metric)
    if not dim_info:
        return
    share_text = pct_fmt(dim_info["share"]) if not pd.isna(dim_info["share"]) else "-"
    add_para(doc, f"图表解读：从{dim_info['dim']}维度看，“{dim_info['top_name']}”贡献最高，合计为{money_fmt(dim_info['top_value'])}，约占总量{share_text}。该结果说明该维度项是当前{metric}的主要来源，但仍需结合记录数、折扣、利润等指标判断其贡献质量。", indent=True)


def _add_risk_interpretation(doc, anomaly_df, main_metric):
    risk = _risk_profile(anomaly_df)
    if risk["unit_n"] == 0:
        return
    add_para(doc, f"图表解读：异常单元占全部经营单元约{pct_fmt(risk['rate'])}，其中高风险单元{risk['high_n']}个。高风险并不等同于一定存在错误，而是说明该对象在{main_metric}、相关数值指标或业务规则上偏离较明显，应进入优先复核清单。", indent=True)


def _add_risk_review_section(doc, anomaly_df, main_metric, dimensions, section_title="重点异常核查清单"):
    if anomaly_df is None or len(anomaly_df) == 0 or "是否经营异常" not in anomaly_df.columns:
        add_para(doc, "当前数据暂未形成稳定的异常核查清单。", indent=True)
        return
    abnormal = anomaly_df[anomaly_df["是否经营异常"]].copy()
    if len(abnormal) == 0:
        add_para(doc, "当前未识别出明显异常经营单元。建议继续关注主指标波动和关键维度结构变化。", indent=True)
        return
    add_heading(doc, section_title, 2)
    review = _build_risk_review_table(anomaly_df, main_metric, dimensions, max_rows=5)
    if len(review):
        add_para(doc, "下表用于回答“优先查谁、为什么查、下一步查什么”。相较于仅展示风险等级和分数，核查清单进一步补充了经营对象、关键指标和建议动作。", indent=True)
        add_dataframe_to_docx(doc, review, max_rows=5)
    add_para(doc, "管理含义：高风险对象通常不是由单一指标造成，而是由高值偏离、多指标同步偏离、折扣/利润等业务规则风险叠加形成。应优先复核原始业务单据、折扣政策、成本核算和审批记录。", indent=True)

def generate_ai_report_docx(ai_text, df, main_metric, dimensions, date_col, anomaly_df, focus_list):
    if not DOCX_AVAILABLE:
        raise RuntimeError("未安装 python-docx，请先安装：python -m pip install python-docx")
    focus_list = clean_focus_list_for_report(focus_list, date_col)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("智策经营——AI增强版管理层决策简报")
    set_run_font(r, "黑体", 20, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向经营管理者的指标表现、结构贡献、风险核查与AI管理建议")
    set_run_font(r, "宋体", 11, False)
    r.font.color.rgb = RGBColor(90, 90, 90)

    ai_paras, ai_tables = extract_clean_ai_blocks(ai_text)

    _add_key_conclusion_box(doc, df, main_metric, dimensions, date_col, anomaly_df)

    add_heading(doc, "二、分析口径与数据概况", 1)
    add_heading(doc, "2.1 数据范围与分析口径", 2)
    add_para(doc, f"本报告基于当前上传数据自动生成。清洗后数据共{len(df):,}条记录，主分析指标为“{main_metric}”。系统仅根据上传数据中实际存在的字段开展分析，不引用未上传或不存在的指标。", indent=True)
    if focus_list:
        add_para(doc, f"本次简报关注重点包括：{'、'.join(focus_list)}。", indent=True)
    if not date_col:
        add_para(doc, "由于当前未选择有效日期字段，系统不生成时间趋势结论，改用结构分布和多维对比方式观察经营表现。", indent=True)

    add_heading(doc, "2.2 核心指标概览", 2)
    profile = _report_metric_profile(df, main_metric)
    risk = _risk_profile(anomaly_df)
    summary_rows = pd.DataFrame([
        {"指标": f"{main_metric}合计", "数值": money_fmt(profile["total"]), "管理含义": "反映总体经营规模"},
        {"指标": f"{main_metric}均值", "数值": money_fmt(profile["mean"]), "管理含义": "反映单条记录的一般水平"},
        {"指标": f"{main_metric}中位数", "数值": money_fmt(profile["median"]), "管理含义": "用于判断典型记录水平"},
        {"指标": f"{main_metric}最大值", "数值": money_fmt(profile["max"]), "管理含义": "用于定位高值经营单元"},
        {"指标": "数据记录数", "数值": f"{len(df):,}", "管理含义": "反映分析样本规模"},
        {"指标": "异常经营单元数", "数值": str(risk["abnormal_n"]), "管理含义": "提示优先核查对象数量"},
    ])
    add_dataframe_to_docx(doc, summary_rows, max_rows=10)

    add_heading(doc, "三、经营表现与结构洞察", 1)
    if date_col:
        trend = build_trend(df, date_col, main_metric)
        add_heading(doc, "3.1 时间趋势分析", 2)
        add_para(doc, trend_interpretation(trend, main_metric), indent=True)
        if len(trend) >= 2:
            add_mpl_figure_to_docx(doc, make_mpl_trend_figure(trend, main_metric), f"图1 {main_metric}时间趋势")
            add_para(doc, "图表解读：趋势图用于判断主指标在不同周期的变化方向和阶段性异常点。", indent=True)
    else:
        add_heading(doc, "3.1 结构分布分析", 2)
        add_para(doc, f"当前未选择有效日期字段，因此重点观察{main_metric}在不同记录中的集中程度、离散程度和高低值分布。", indent=True)
        add_mpl_figure_to_docx(doc, make_mpl_distribution_figure(df, main_metric), f"图1 {main_metric}分布情况")
        _add_distribution_interpretation(doc, df, main_metric)

    if dimensions:
        add_heading(doc, "3.2 多维结构贡献", 2)
        add_mpl_figure_to_docx(doc, make_mpl_dimension_bar_figure(df, dimensions[0], main_metric), f"图2 按{dimensions[0]}汇总{main_metric}")
        _add_dimension_interpretation(doc, df, dimensions, main_metric)

    add_heading(doc, "四、重点风险与异常核查", 1)
    if anomaly_df is not None and len(anomaly_df) and '是否经营异常' in anomaly_df.columns:
        add_para(doc, f"系统当前识别出{risk['abnormal_n']}个异常经营单元，其中高风险单元{risk['high_n']}个。风险结果用于提示优先核查对象，并需结合企业真实业务背景进行复核。", indent=True)
        add_mpl_figure_to_docx(doc, make_mpl_risk_figure(anomaly_df), "图3 风险等级分布")
        _add_risk_interpretation(doc, anomaly_df, main_metric)
        _add_risk_review_section(doc, anomaly_df, main_metric, dimensions, section_title="4.1 重点异常核查清单")
    else:
        add_para(doc, "当前数据暂未形成稳定的异常诊断结果。", indent=True)

    add_heading(doc, "五、AI增强解读与管理建议", 1)
    add_heading(doc, "5.1 AI综合判断", 2)
    useful = [p for p in ai_paras if not is_report_meta_line(p)]
    # 只保留管理判断类内容，避免AI再次生成报告标题、数据周期和重复大纲
    useful = [p for p in useful if len(clean_report_sentence(p)) >= 12][:4]
    if useful:
        for p in useful:
            add_para(doc, p, indent=True)
    else:
        add_para(doc, "AI增强解读用于在系统统计结果基础上进一步归纳经营问题和管理动作。当前结果显示，应重点关注高值经营单元、主要贡献维度和异常风险对象的业务真实性与盈利质量。", indent=True)

    add_heading(doc, "5.2 管理行动清单", 2)
    used_table = False
    if ai_tables:
        for _, tb in ai_tables[:1]:
            if tb is not None and len(tb) and tb.shape[1] >= 2:
                add_dataframe_to_docx(doc, tb, max_rows=10)
                used_table = True
                break
    if not used_table:
        action_df = _build_management_action_table(main_metric, dimensions, anomaly_df)
        add_dataframe_to_docx(doc, action_df, max_rows=10)

    add_para(doc, "以上建议用于辅助管理层确定核查优先级。实际决策仍需结合企业业务背景、政策制度、预算目标和原始业务单据进行综合判断。", indent=True)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def anomaly_detection(df, main_metric, numeric_cols, dimensions, date_col=None):
    agg_cols = [c for c in numeric_cols if c in df.columns and c != "" and not is_id_like(c)]
    if len(agg_cols) == 0 or main_metric not in df.columns:
        return pd.DataFrame()

    temp = df.copy()
    group_cols = []
    if date_col:
        temp = periodize(temp, date_col)
        if not temp.empty:
            group_cols.append("_period")
    for d in dimensions[:2]:
        if d in temp.columns:
            group_cols.append(d)

    if group_cols:
        # 不同类型经营指标采用不同聚合口径：
        # 金额/数量/利润等规模指标求和；折扣/比率/单价/时长等水平指标取平均。
        agg_map = {c: aggregation_method_for_metric(c) for c in agg_cols}
        unit = temp.groupby(group_cols, dropna=False).agg(agg_map).reset_index()
        unit["记录数"] = temp.groupby(group_cols, dropna=False).size().values
    else:
        unit = temp[agg_cols].copy()
        unit["记录数"] = 1

    model_cols = [c for c in agg_cols if c in unit.columns and unit[c].notna().sum() > 0]
    if len(model_cols) == 0:
        return pd.DataFrame()

    X_df = unit[model_cols].replace([np.inf, -np.inf], np.nan)
    X_df = X_df.fillna(X_df.median(numeric_only=True))

    if len(unit) >= 8 and len(model_cols) >= 1:
        scaler = StandardScaler()
        X = scaler.fit_transform(X_df)
        contamination = min(0.12, max(0.04, 8 / max(len(unit), 1)))
        model = IsolationForest(n_estimators=160, contamination=contamination, random_state=42)
        labels = model.fit_predict(X)
        raw = -model.decision_function(X)
        unit["是否AI异常"] = labels == -1
        unit["异常得分"] = 0 if raw.max() == raw.min() else (raw - raw.min()) / (raw.max() - raw.min()) * 100
    else:
        unit["是否AI异常"] = False
        unit["异常得分"] = 0.0

    q95 = X_df.quantile(0.95)
    q05 = X_df.quantile(0.05)

    evidence_list = []
    main_component = []
    combo_component = []
    rule_component = []

    for _, row in unit.iterrows():
        evidences = []
        mc = cc = rc = 0

        if main_metric in unit.columns:
            val = row[main_metric]
            pct_rank = (unit[main_metric] <= val).mean()
            if pct_rank >= 0.95:
                evidences.append(f"{main_metric}处于高位区间，高于约{pct_rank:.0%}的经营单元")
                mc = 30
            elif pct_rank <= 0.05:
                evidences.append(f"{main_metric}处于低位区间，低于约{(1-pct_rank):.0%}的经营单元")
                mc = 20

        extreme_cols = []
        for c in model_cols:
            val = row[c]
            if val >= q95[c] or val <= q05[c]:
                extreme_cols.append(c)
        if len(extreme_cols) >= 2:
            evidences.append(f"多项指标同时处于极端区间：{', '.join(extreme_cols[:4])}")
            cc = min(10 * len(extreme_cols), 35)
        elif len(extreme_cols) == 1:
            evidences.append(f"{extreme_cols[0]}处于极端区间")
            cc = 12

        for c in model_cols:
            cname = c.lower()
            val = row[c]
            if ("profit" in cname or "利润" in cname) and val < 0:
                evidences.append(f"{c}为负，存在收益或盈利风险")
                rc += 20
            if is_rate_like(c) and (val > 1.2 or val < -0.2):
                evidences.append(f"{c}超出常见比例范围，需要复核口径")
                rc += 10
            if ("cost" in cname or "费用" in cname or "成本" in cname or "expense" in cname) and val >= q95[c]:
                evidences.append(f"{c}处于高位区间，提示资源投入或支出压力较高")
                rc += 12

        if not evidences and row.get("是否AI异常", False):
            evidences.append("多指标组合与整体样本差异较大，被模型识别为综合异常")
        if not evidences:
            evidences.append("未识别出显著异常依据")

        evidence_list.append("；".join(evidences))
        main_component.append(mc)
        combo_component.append(cc)
        rule_component.append(min(rc, 25))

    unit["异常依据"] = evidence_list
    unit["主指标偏离"] = main_component
    unit["多指标组合偏离"] = combo_component
    unit["业务规则风险"] = rule_component
    unit["模型异常贡献"] = np.where(unit["是否AI异常"], np.minimum(unit["异常得分"] * 0.35, 35), np.minimum(unit["异常得分"] * 0.15, 15))
    unit["风险得分"] = (unit["主指标偏离"] + unit["多指标组合偏离"] + unit["业务规则风险"] + unit["模型异常贡献"]).clip(0, 100).round(1)

    def level(score):
        if score >= 70:
            return "高风险"
        if score >= 45:
            return "中风险"
        if score > 0:
            return "低风险"
        return "正常"

    unit["风险等级"] = unit["风险得分"].apply(level)
    unit["是否经营异常"] = unit["风险得分"] >= 45
    return unit.sort_values("风险得分", ascending=False)


def radar_chart(comp):
    theta = comp["风险来源"].tolist()
    r = comp["得分"].tolist()
    if theta:
        theta = theta + [theta[0]]
        r = r + [r[0]]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=r, theta=theta, fill="toself", name="风险组成"))
    fig.update_layout(
        title="风险雷达图",
        polar=dict(radialaxis=dict(visible=True, range=[0, max(40, max(r) if r else 40)])),
        showlegend=False
    )
    st.plotly_chart(chart_layout(fig, 380), use_container_width=True)


# ============================================================
# 4. 问数助手
# ============================================================

def get_secret_value(name):
    try:
        if name in st.secrets:
            return st.secrets[name]
    except Exception:
        pass
    return os.getenv(name, "")


def get_llm_configs():
    return {
        "通义千问": {
            "api_key": get_secret_value("QWEN_API_KEY"),
            "base_url": get_secret_value("QWEN_BASE_URL") or "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
            "model": get_secret_value("QWEN_MODEL") or "qwen-plus"
        },
        "DeepSeek": {
            "api_key": get_secret_value("DEEPSEEK_API_KEY"),
            "base_url": get_secret_value("DEEPSEEK_BASE_URL") or "https://api.deepseek.com/chat/completions",
            "model": get_secret_value("DEEPSEEK_MODEL") or "deepseek-chat"
        },
        "智谱GLM": {
            "api_key": get_secret_value("ZHIPU_API_KEY"),
            "base_url": get_secret_value("ZHIPU_BASE_URL") or "https://open.bigmodel.cn/api/paas/v4/chat/completions",
            "model": get_secret_value("ZHIPU_MODEL") or "glm-4-flash"
        }
    }


def call_llm(provider_name, system_prompt, user_prompt, temperature=0.1):
    cfg = get_llm_configs()[provider_name]
    if not cfg["api_key"]:
        raise ValueError(f"{provider_name} 未配置 API Key。请在 .streamlit/secrets.toml 中配置。")
    headers = {"Authorization": f"Bearer {cfg['api_key']}", "Content-Type": "application/json"}
    payload = {
        "model": cfg["model"],
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": temperature
    }
    start = time.time()
    resp = requests.post(cfg["base_url"], headers=headers, json=payload, timeout=60)
    elapsed = time.time() - start
    if resp.status_code != 200:
        raise RuntimeError(f"API 请求失败：{resp.status_code}｜{resp.text[:500]}")
    data = resp.json()
    return data["choices"][0]["message"]["content"], elapsed


def to_sqlite_df(df):
    sql_df = df.copy()
    sql_df.columns = [normalize_column_name(c) for c in sql_df.columns]
    for col in sql_df.columns:
        if pd.api.types.is_datetime64_any_dtype(sql_df[col]):
            sql_df[col] = sql_df[col].dt.strftime("%Y-%m-%d")
    return sql_df


def build_schema_text(df, meta=None):
    lines = []
    for col in df.columns:
        dtype = str(df[col].dtype)
        samples = df[col].dropna().astype(str).head(3).tolist()
        meaning = ""
        if meta is not None and col in meta["字段名"].values:
            meaning = meta.loc[meta["字段名"] == col, "字段含义解释"].iloc[0]
        lines.append(f"- {col}｜类型：{dtype}｜含义：{meaning}｜示例：{samples}")
    return "\n".join(lines)


def extract_sql(text):
    if not text:
        return ""
    text = text.strip()
    m = re.search(r"```sql(.*?)```", text, flags=re.S | re.I)
    if m:
        return m.group(1).strip()
    m = re.search(r"```(.*?)```", text, flags=re.S)
    if m:
        return m.group(1).strip()
    m = re.search(r"(SELECT\s+.*)", text, flags=re.S | re.I)
    if m:
        sql = m.group(1).strip()
        if ";" in sql:
            sql = sql.split(";")[0] + ";"
        return sql
    return text


def is_safe_select_sql(sql):
    if not sql:
        return False, "SQL 为空"
    low = sql.strip().strip(";").lower()
    if not low.startswith("select"):
        return False, "只允许 SELECT 查询"
    banned = ["insert", "update", "delete", "drop", "alter", "create", "truncate", "replace", "attach", "detach", "pragma"]
    if any(re.search(rf"\b{b}\b", low) for b in banned):
        return False, "检测到非查询类 SQL，已拦截"
    return True, ""


def score_llm_query(ok, result_df, sql, explanation, elapsed, columns):
    score = 0
    detail = {}

    detail["SQL可执行性"] = 40 if ok else 0
    score += detail["SQL可执行性"]

    if ok and isinstance(result_df, pd.DataFrame) and len(result_df) > 0:
        detail["查询结果有效性"] = 20
    elif ok:
        detail["查询结果有效性"] = 8
    else:
        detail["查询结果有效性"] = 0
    score += detail["查询结果有效性"]

    used_cols = [c for c in columns if str(c).lower() in str(sql).lower()]
    detail["字段匹配度"] = 15 if used_cols else (5 if ok else 0)
    score += detail["字段匹配度"]

    detail["结果解释质量"] = 15 if explanation and len(str(explanation)) > 40 else (6 if explanation else 0)
    score += detail["结果解释质量"]

    if pd.isna(elapsed):
        speed = 0
    elif elapsed <= 3:
        speed = 10
    elif elapsed <= 8:
        speed = 8
    elif elapsed <= 15:
        speed = 5
    else:
        speed = 2
    detail["响应速度"] = speed
    score += speed

    return min(int(score), 100), detail



def _find_cols_by_keywords(columns, keywords):
    hits = []
    for c in columns:
        s = str(c).lower()
        if any(k.lower() in s for k in keywords):
            hits.append(c)
    return hits


def infer_question_caliber(question, sql, result_df, df, meta):
    """根据用户问题、SQL和当前字段生成问数口径说明。"""
    q = str(question or "")
    cols = list(df.columns)
    rows = []

    used_cols = []
    for c in cols:
        if f'"{c}"' in str(sql) or f'`{c}`' in str(sql) or re.search(rf'(?<![\w\u4e00-\u9fff]){re.escape(str(c))}(?![\w\u4e00-\u9fff])', str(sql), flags=re.I):
            used_cols.append(c)

    metric_guess = ""
    if any(k in q for k in ["销售量", "销量", "销售数量", "数量最高", "数量最多"]):
        cand = _find_cols_by_keywords(cols, ["数量", "qty", "quantity", "count"])
        metric_guess = f"销售量/销量 → 优先理解为“{cand[0]}”的合计" if cand else "销售量/销量 → 当前数据中未明显识别到数量类字段，需结合SQL核对"
    elif any(k in q for k in ["销售额", "收入", "金额", "利润", "成本", "费用"]):
        cand = [c for c in cols if any(k in str(c) for k in ["销售", "收入", "金额", "利润", "成本", "费用"])]
        metric_guess = f"经营指标 → 优先匹配为“{cand[0]}”等金额/利润类字段" if cand else "经营指标 → 按SQL中使用的数值字段解释"
    elif used_cols:
        metric_guess = f"经营指标 → SQL中实际使用字段：{', '.join(map(str, used_cols[:4]))}"
    else:
        metric_guess = "经营指标 → 按大模型生成SQL中的字段执行"

    if any(k in q for k in ["产品", "商品", "sku", "SKU"]):
        prod_cols = _find_cols_by_keywords(cols, ["产品", "商品", "sku", "SKU"])
        obj_guess = f"产品/商品 → 优先匹配为“{prod_cols[0]}”" if prod_cols else "产品/商品 → 当前数据未明显识别到产品字段，需检查SQL分组字段"
    elif any(k in q for k in ["客户", "地区", "部门", "类别", "类型", "渠道"]):
        dim_hits = [c for c in cols if any(k in str(c) for k in ["客户", "地区", "部门", "类别", "类型", "渠道"])]
        obj_guess = f"分析对象 → 优先匹配为“{dim_hits[0]}”等维度字段" if dim_hits else "分析对象 → 按SQL中的GROUP BY字段解释"
    else:
        obj_guess = "分析对象 → 按SQL中的分组字段或筛选条件确定"

    if any(k in q.lower() for k in ["最高", "最大", "最多", "top", "前10", "前十", "排名"]):
        sort_guess = "排序方式 → 按目标指标降序排序，返回排名靠前结果"
    elif any(k in q for k in ["最低", "最小", "最少"]):
        sort_guess = "排序方式 → 按目标指标升序排序，返回排名靠前结果"
    else:
        sort_guess = "排序方式 → 未明显指定排名，按SQL生成逻辑执行"

    rows.append({"解析项": "指标口径", "系统理解": metric_guess})
    rows.append({"解析项": "分析对象", "系统理解": obj_guess})
    rows.append({"解析项": "排序/筛选", "系统理解": sort_guess})
    rows.append({"解析项": "查询范围", "系统理解": "当前上传并清洗后的数据表 data"})
    if used_cols:
        rows.append({"解析项": "SQL实际字段", "系统理解": "、".join(map(str, used_cols[:8]))})
    return pd.DataFrame(rows)


def build_query_validation(pack, question, df):
    sql = str(pack.get("SQL", ""))
    result = pack.get("结果", pd.DataFrame())
    ok = bool(pack.get("是否成功", False))
    used_cols = []
    for c in df.columns:
        if f'"{c}"' in sql or f'`{c}`' in sql or re.search(rf'(?<![\w\u4e00-\u9fff]){re.escape(str(c))}(?![\w\u4e00-\u9fff])', sql, flags=re.I):
            used_cols.append(c)

    agg = []
    low_sql = sql.lower()
    if "sum(" in low_sql:
        agg.append("包含 SUM 求和聚合")
    if "avg(" in low_sql or "mean(" in low_sql:
        agg.append("包含 AVG 均值聚合")
    if "count(" in low_sql:
        agg.append("包含 COUNT 计数")
    if "group by" in low_sql:
        agg.append("包含 GROUP BY 分组")
    if "order by" in low_sql:
        agg.append("包含 ORDER BY 排序")

    rows = [
        {"校验项": "SQL可执行性", "校验结果": "通过" if ok else "未通过"},
        {"校验项": "字段匹配", "校验结果": "、".join(map(str, used_cols)) if used_cols else "未识别到明确字段匹配，需人工复核"},
        {"校验项": "聚合/排序逻辑", "校验结果": "；".join(agg) if agg else "未识别到明显聚合或排序逻辑"},
        {"校验项": "结果有效性", "校验结果": f"返回 {len(result)} 行结果" if isinstance(result, pd.DataFrame) else "无结果表"},
        {"校验项": "可信度提示", "校验结果": "结果来自SQL执行；若问题口径存在歧义，请优先核对指标口径和分组字段"}
    ]
    return pd.DataFrame(rows)


def extract_direct_answer_from_explanation(explanation):
    text = clean_md_inline(explanation)
    if not text:
        return "模型已返回查询结果，但未生成直接回答。请结合SQL结果表进行判断。"
    # 优先提取“直接回答/查询结论”后的内容
    m = re.search(r"(?:直接回答|查询结论|结论)[:：]\s*(.*?)(?:经营含义|后续建议|建议[:：]|$)", text, flags=re.S)
    if m:
        ans = m.group(1).strip(" ：；。")
        if ans:
            return ans[:260] + ("..." if len(ans) > 260 else "")
    # 兜底取第一句
    parts = re.split(r"[。；\n]", text)
    for p in parts:
        p = p.strip()
        if len(p) >= 8:
            return p[:260] + ("..." if len(p) > 260 else "")
    return text[:260] + ("..." if len(text) > 260 else "")


def render_llm_query_pack(pack, question, df, meta):
    """问数结果可信度闭环展示。"""
    provider = pack.get("模型", "")
    st.markdown(f"## {provider}")

    if pack.get("是否成功"):
        direct = pack.get("直接回答") or extract_direct_answer_from_explanation(pack.get("解释", ""))
        st.markdown(f"""
        <div class="section-card">
            <h3>AI直接回答</h3>
            <p class="small-text">{direct}</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("#### 问题解析口径")
        caliber_df = pack.get("问题解析口径")
        if not isinstance(caliber_df, pd.DataFrame):
            caliber_df = infer_question_caliber(question, pack.get("SQL", ""), pack.get("结果", pd.DataFrame()), df, meta)
        st.dataframe(caliber_df, use_container_width=True, hide_index=True)

        st.markdown("#### SQL 查询语句")
        st.code(pack.get("SQL", ""), language="sql")

        st.success(f"执行成功｜响应时间：{pack.get('响应时间', np.nan):.2f}秒｜综合评分：{pack.get('综合评分', 0)}")
        st.markdown("#### 数据库执行结果")
        st.dataframe(pack.get("结果", pd.DataFrame()), use_container_width=True)

        st.markdown("#### 结果解释")
        st.info(pack.get("解释", ""))

        st.markdown("#### 系统校验与评分")
        validation_df = pack.get("系统校验")
        if not isinstance(validation_df, pd.DataFrame):
            validation_df = build_query_validation(pack, question, df)
        st.dataframe(validation_df, use_container_width=True, hide_index=True)

        with st.expander("查看评分明细"):
            st.json(pack.get("评分明细", {}))
    else:
        if pack.get("SQL"):
            st.markdown("#### SQL 查询语句")
            st.code(pack.get("SQL", ""), language="sql")
        st.error(f"执行失败：{pack.get('错误信息', '')}")
        with st.expander("查看评分明细"):
            st.json(pack.get("评分明细", {}))


def run_llm_sql_question(provider, question, df, meta):
    sql_df = to_sqlite_df(df)
    con = sqlite3.connect(":memory:")
    sql_df.to_sql("data", con, index=False, if_exists="replace")
    schema = build_schema_text(sql_df, meta=None)

    system_prompt = """
你是一个严谨的数据分析 SQL 助手。你的任务是把用户的自然语言问题转换为 SQLite SELECT 查询语句。
要求：
1. 只输出 SQL，不要解释；
2. 表名固定为 data；
3. 只能使用 SELECT；
4. 字段名必须严格来自字段信息，中文或特殊字段名请使用双引号；
5. 尽量使用 LIMIT 20 控制输出；
6. 若问题涉及排名、最高、最低，请使用 ORDER BY；
7. 若问题涉及分组统计，请使用 GROUP BY；
8. 若问题涉及“销售量/销量/销售数量”，优先寻找数量类字段并对其求和；
9. 若问题涉及“产品/商品/SKU”，优先寻找产品、商品、SKU相关字段作为分组对象。
"""
    user_prompt = f"""字段信息：
{schema}

用户问题：{question}

请生成 SQLite SELECT 查询语句。"""
    content, elapsed = call_llm(provider, system_prompt, user_prompt, temperature=0)
    sql = extract_sql(content)

    safe, reason = is_safe_select_sql(sql)
    if not safe:
        score, detail = score_llm_query(False, pd.DataFrame(), sql, "", elapsed, sql_df.columns)
        pack = {"模型": provider, "SQL": sql, "是否成功": False, "错误信息": reason, "结果": pd.DataFrame(), "解释": "", "直接回答": "", "响应时间": elapsed, "综合评分": score, "评分明细": detail}
        pack["问题解析口径"] = infer_question_caliber(question, sql, pd.DataFrame(), df, meta)
        pack["系统校验"] = build_query_validation(pack, question, df)
        return pack

    try:
        result = pd.read_sql_query(sql, con)
        ok, error = True, ""
    except Exception as e:
        result = pd.DataFrame()
        ok, error = False, str(e)

    explanation = ""
    direct_answer = ""
    if ok:
        try:
            exp_system = "你是一名经营数据分析顾问。请根据SQL查询结果，先直接回答用户问题，再解释经营含义和后续建议。回答必须基于查询结果，不得编造结果表中不存在的字段或数值。"
            result_text = result.head(20).to_markdown(index=False)
            exp_user = f"""用户问题：{question}
SQL：{sql}
查询结果：
{result_text}

请按以下格式输出：
直接回答：...
结果解释：...
经营含义：...
后续建议：..."""
            explanation, _ = call_llm(provider, exp_system, exp_user, temperature=0.2)
            direct_answer = extract_direct_answer_from_explanation(explanation)
        except Exception as e:
            explanation = f"SQL 已执行成功，但结果解释生成失败：{e}"
            direct_answer = "SQL 已执行成功，请以数据库执行结果表为准。"

    score, detail = score_llm_query(ok, result, sql, explanation, elapsed, sql_df.columns)
    pack = {"模型": provider, "SQL": sql, "是否成功": ok, "错误信息": error, "结果": result, "解释": explanation, "直接回答": direct_answer, "响应时间": elapsed, "综合评分": score, "评分明细": detail}
    pack["问题解析口径"] = infer_question_caliber(question, sql, result, df, meta)
    pack["系统校验"] = build_query_validation(pack, question, df)
    return pack

def generate_dynamic_questions(main_metric, dimensions, numeric_cols, date_col):
    qs = []
    d1 = dimensions[0] if dimensions else None
    d2 = dimensions[1] if len(dimensions) > 1 else None
    other = next((c for c in numeric_cols if c != main_metric and not is_id_like(c)), None)

    if d1:
        qs.append(("排名类", f"按 {d1} 统计 {main_metric} 最高的前10项"))
        qs.append(("对比类", f"比较不同 {d1} 的 {main_metric} 均值差异"))
        qs.append(("结构类", f"各 {d1} 的 {main_metric} 占比分别是多少"))
    if d2:
        qs.append(("交叉分析类", f"按 {d1} 和 {d2} 交叉统计 {main_metric}"))
    if date_col:
        qs.append(("趋势类", f"按月份统计 {main_metric} 的变化趋势"))
        qs.append(("波动类", f"找出 {main_metric} 变化最大的月份"))
    if other:
        qs.append(("关系类", f"分析 {other} 与 {main_metric} 是否存在关系"))
    qs.append(("异常类", f"查询 {main_metric} 异常偏高或偏低的记录"))
    qs.append(("风险类", "当前风险最高的经营单元有哪些"))
    return qs


# ============================================================
# 5. Word 简报
# ============================================================

def set_run_font(run, font="宋体", size=11, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "黑体", 15 if level == 1 else 13, True)
    return p


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(clean_report_sentence(text))
    set_run_font(run, "宋体", 11, False)
    return p
def _pick_mpl_chinese_font():
    """选择 Matplotlib 可用中文字体，返回字体名；找不到时返回空字符串。"""
    try:
        from matplotlib import font_manager
        candidates = [
            r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simsun.ttc",
            "/System/Library/Fonts/PingFang.ttc", "/System/Library/Fonts/STHeiti Light.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.otf",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
        ]
        for fp in candidates:
            if os.path.exists(fp):
                try:
                    font_manager.fontManager.addfont(fp)
                    return font_manager.FontProperties(fname=fp).get_name()
                except Exception:
                    continue
        for f in font_manager.fontManager.ttflist:
            name = f.name.lower()
            if any(k.lower() in name for k in ["Microsoft YaHei", "SimHei", "SimSun", "PingFang", "Noto Sans CJK", "WenQuanYi"]):
                return f.name
    except Exception:
        pass
    return ""
_ZC_HAS_CHINESE_FONT = False

def _setup_mpl_style():
    global _ZC_HAS_CHINESE_FONT
    font_name = _pick_mpl_chinese_font()
    if font_name:
        plt.rcParams["font.sans-serif"] = [font_name, "Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
        _ZC_HAS_CHINESE_FONT = True
    else:
        plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
        _ZC_HAS_CHINESE_FONT = False
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"


def _mpl_has_chinese_font():
    return bool(_ZC_HAS_CHINESE_FONT)
def _short_label(x, max_len=14):
    s = str(x)
    return s if len(s) <= max_len else s[:max_len] + "…"


def _mpl_money_axis(x, pos=None):
    """Matplotlib 图内尽量使用 ASCII 数值后缀，避免云端环境缺少中文字体时出现方框。"""
    try:
        x = float(x)
    except Exception:
        return str(x)
    sign = "-" if x < 0 else ""
    ax = abs(x)
    if ax >= 1_000_000_000:
        return f"{sign}{ax/1_000_000_000:.1f}B"
    if ax >= 1_000_000:
        return f"{sign}{ax/1_000_000:.1f}M"
    if ax >= 1_000:
        return f"{sign}{ax/1_000:.0f}K"
    return f"{x:.0f}"


def _mpl_value_label(x):
    """Matplotlib 图内数据标签，避免使用“万/亿”等中文单位。"""
    try:
        x = float(x)
    except Exception:
        return str(x)
    sign = "-" if x < 0 else ""
    ax = abs(x)
    if ax >= 1_000_000_000:
        return f"{sign}{ax/1_000_000_000:.2f}B"
    if ax >= 1_000_000:
        return f"{sign}{ax/1_000_000:.2f}M"
    if ax >= 1_000:
        return f"{sign}{ax/1_000:.1f}K"
    return f"{x:.2f}"

def add_mpl_figure_to_docx(doc, fig, caption):
    """把 Matplotlib 图表以 PNG 插入 Word；图题统一放在图下方，符合正式报告习惯。"""
    try:
        img = BytesIO()
        fig.savefig(img, format="png", dpi=190, bbox_inches="tight", facecolor="white")
        mapping = getattr(fig, "_zc_mapping", None)
        plt.close(fig)
        img.seek(0)

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img, width=Inches(6.2))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(8)
        run = p_cap.add_run(caption)
        set_run_font(run, "宋体", 10, True)

        if mapping:
            map_df = pd.DataFrame(mapping)
            if len(map_df):
                p2 = doc.add_paragraph()
                run2 = p2.add_run("图中编码说明：")
                set_run_font(run2, "黑体", 10, True)
                add_dataframe_to_docx(doc, map_df, max_rows=20)
        return True
    except Exception as e:
        try:
            plt.close(fig)
        except Exception:
            pass
        add_para(doc, f"图表“{caption}”生成失败：{e}。", indent=False)
        return False

def make_mpl_trend_figure(trend, metric):
    _setup_mpl_style()
    has_cn = _mpl_has_chinese_font()
    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    plot = trend.copy()
    plot[metric] = pd.to_numeric(plot[metric], errors="coerce")
    plot = plot.dropna(subset=["期间", metric])
    xs = list(range(len(plot)))
    ys = plot[metric].astype(float).tolist()
    ax.plot(xs, ys, linewidth=2.8, marker="o", markersize=4.8)
    ax.fill_between(xs, ys, alpha=0.10)
    if len(xs) > 0:
        max_i = int(np.nanargmax(ys)); min_i = int(np.nanargmin(ys))
        ax.scatter([max_i], [ys[max_i]], s=78, zorder=4)
        ax.scatter([min_i], [ys[min_i]], s=78, zorder=4)
        ax.annotate(f"Max {_mpl_value_label(ys[max_i])}", (max_i, ys[max_i]), xytext=(0, 12), textcoords="offset points", ha="center", fontsize=9)
        ax.annotate(f"Min {_mpl_value_label(ys[min_i])}", (min_i, ys[min_i]), xytext=(0, -18), textcoords="offset points", ha="center", fontsize=9)
    periods = plot["期间"].astype(str).tolist()
    tick_idx = list(range(len(periods)))
    if len(periods) > 8:
        step = max(1, int(np.ceil(len(periods) / 8)))
        tick_idx = list(range(0, len(periods), step))
        if len(periods)-1 not in tick_idx:
            tick_idx.append(len(periods)-1)
    ax.set_xticks(tick_idx)
    ax.set_xticklabels([periods[i] for i in tick_idx], rotation=0, fontsize=9)
    ax.yaxis.set_major_formatter(FuncFormatter(_mpl_money_axis))
    ax.set_title(f"{metric}时间趋势" if has_cn else "Trend of Key Metric", fontsize=14, fontweight="bold", loc="left")
    ax.set_xlabel("期间" if has_cn else "Period")
    ax.set_ylabel(metric if has_cn else "Value")
    ax.grid(axis="y", alpha=0.25)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return fig
def make_mpl_dimension_bar_figure(df, dim, metric, topn=10):
    _setup_mpl_style()
    has_cn = _mpl_has_chinese_font()
    g = dimension_summary(df, dim, metric).head(topn).copy()
    g[f"{metric}合计"] = pd.to_numeric(g[f"{metric}合计"], errors="coerce").fillna(0)
    g = g.sort_values(f"{metric}合计", ascending=True)
    fig, ax = plt.subplots(figsize=(8.8, max(4.2, 0.42 * len(g) + 1.2)))
    raw_labels = g[dim].astype(str).tolist()
    if has_cn:
        y_labels = [_short_label(x, 12) for x in raw_labels]
    else:
        y_labels = [f"C{i+1}" for i in range(len(raw_labels))]
        fig._zc_mapping = [{"图中编码": code, "对应维度项": raw} for code, raw in zip(y_labels, raw_labels)]
    vals = g[f"{metric}合计"].astype(float).tolist()
    bars = ax.barh(y_labels, vals, height=0.62)
    ax.xaxis.set_major_formatter(FuncFormatter(_mpl_money_axis))
    ax.set_title(f"按{dim}汇总{metric}" if has_cn else "Dimension Summary", fontsize=14, fontweight="bold", loc="left")
    ax.set_xlabel(f"{metric}合计" if has_cn else "Total Value")
    ax.set_ylabel(dim if has_cn else "Category Code")
    ax.grid(axis="x", alpha=0.22)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    max_val = max([abs(v) for v in vals] + [1])
    for bar, val in zip(bars, vals):
        ax.text(val + max_val * 0.015, bar.get_y() + bar.get_height()/2, _mpl_value_label(val), va="center", fontsize=9)
    fig.tight_layout()
    return fig
def make_mpl_distribution_figure(df, metric):
    _setup_mpl_style()
    has_cn = _mpl_has_chinese_font()
    s = to_finite_numeric_series(df, metric)
    fig, ax = plt.subplots(figsize=(8.8, 4.4))
    if len(s) > 0:
        ax.hist(s.astype(float).tolist(), bins=min(30, max(8, int(np.sqrt(len(s))))), alpha=0.86)
        ax.axvline(float(s.mean()), linestyle="--", linewidth=2, label=("均值" if has_cn else "Mean"))
        ax.legend(fontsize=9)
    ax.xaxis.set_major_formatter(FuncFormatter(_mpl_money_axis))
    ax.set_title(f"{metric}分布情况" if has_cn else "Metric Distribution", fontsize=14, fontweight="bold", loc="left")
    ax.set_xlabel(metric if has_cn else "Value")
    ax.set_ylabel("记录数" if has_cn else "Count")
    ax.grid(axis="y", alpha=0.22)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return fig
def make_mpl_risk_figure(anomaly_df):
    _setup_mpl_style()
    has_cn = _mpl_has_chinese_font()
    fig, ax = plt.subplots(figsize=(8.0, 4.2))
    mapping = []
    if anomaly_df is not None and len(anomaly_df) and "风险等级" in anomaly_df.columns:
        order = ["高风险", "中风险", "低风险", "正常"]
        counts = anomaly_df["风险等级"].value_counts()
        labels = [x for x in order if x in counts.index] + [x for x in counts.index if x not in order]
        vals = [int(counts.get(x, 0)) for x in labels]
        if has_cn:
            x_labels = labels
        else:
            x_labels = [f"R{i+1}" for i in range(len(labels))]
            mapping = [{"图中编码": code, "风险等级": label} for code, label in zip(x_labels, labels)]
        bars = ax.bar(x_labels, vals, width=0.55)
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width()/2, val + max(vals + [1])*0.02, str(val), ha="center", fontsize=10)
    ax.set_title("风险等级分布" if has_cn else "Risk Level Distribution", fontsize=14, fontweight="bold", loc="left")
    ax.set_xlabel("风险等级" if has_cn else "Risk Code")
    ax.set_ylabel("数量" if has_cn else "Count")
    ax.grid(axis="y", alpha=0.22)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    if mapping:
        fig._zc_mapping = mapping
    fig.tight_layout()
    return fig
def add_core_visuals_to_docx(doc, df, main_metric, dimensions, date_col, anomaly_df):
    """在 Word 简报中插入核心可视化：趋势/分布、维度结构、风险分布。"""
    inserted = 0
    if date_col:
        trend = build_trend(df, date_col, main_metric)
        if trend is not None and len(trend) >= 2:
            add_mpl_figure_to_docx(doc, make_mpl_trend_figure(trend, main_metric), f"图1 {main_metric}时间趋势")
            inserted += 1
    if inserted == 0:
        add_mpl_figure_to_docx(doc, make_mpl_distribution_figure(df, main_metric), f"图1 {main_metric}分布情况")
        inserted += 1
    if dimensions:
        add_mpl_figure_to_docx(doc, make_mpl_dimension_bar_figure(df, dimensions[0], main_metric), f"图2 按{dimensions[0]}汇总{main_metric}")
        inserted += 1
    if anomaly_df is not None and len(anomaly_df):
        add_mpl_figure_to_docx(doc, make_mpl_risk_figure(anomaly_df), "图3 风险等级分布")
        inserted += 1
    return inserted

def generate_report_docx(df, main_metric, dimensions, date_col, anomaly_df, focus_list):
    if not DOCX_AVAILABLE:
        raise RuntimeError("未安装 python-docx，请先安装：python -m pip install python-docx")
    focus_list = clean_focus_list_for_report(focus_list, date_col)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("智策经营——管理层决策简报")
    set_run_font(r, "黑体", 20, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向经营管理者的指标表现、结构贡献与风险核查报告")
    set_run_font(r, "宋体", 11, False)
    r.font.color.rgb = RGBColor(90, 90, 90)

    _add_key_conclusion_box(doc, df, main_metric, dimensions, date_col, anomaly_df)

    add_heading(doc, "二、分析口径与数据概况", 1)
    add_heading(doc, "2.1 数据范围与分析口径", 2)
    add_para(doc, f"本报告基于用户上传的数据自动生成。当前清洗后数据共包含{len(df):,}条记录，主分析指标为“{main_metric}”。系统仅根据上传数据中实际存在的字段开展分析，不引用未上传或不存在的指标。", indent=True)
    if focus_list:
        add_para(doc, f"本次简报关注重点包括：{'、'.join(focus_list)}。", indent=True)
    if not date_col:
        add_para(doc, "当前未选择有效日期字段，因此不生成时间趋势结论，系统转为结构分布、多维对比和异常核查分析。", indent=True)

    add_heading(doc, "2.2 核心指标概览", 2)
    profile = _report_metric_profile(df, main_metric)
    abnormal_n = int((anomaly_df['是否经营异常']).sum()) if anomaly_df is not None and len(anomaly_df) and '是否经营异常' in anomaly_df.columns else 0
    summary_rows = pd.DataFrame([
        {"指标": f"{main_metric}合计", "数值": money_fmt(profile["total"]), "管理含义": "反映总体经营规模"},
        {"指标": f"{main_metric}均值", "数值": money_fmt(profile["mean"]), "管理含义": "反映单条记录的一般水平"},
        {"指标": f"{main_metric}中位数", "数值": money_fmt(profile["median"]), "管理含义": "用于观察典型记录水平，避免被极端值影响"},
        {"指标": f"{main_metric}最大值", "数值": money_fmt(profile["max"]), "管理含义": "用于定位高值经营单元"},
        {"指标": "数据记录数", "数值": f"{len(df):,}", "管理含义": "反映分析样本规模"},
        {"指标": "异常经营单元数", "数值": str(abnormal_n), "管理含义": "提示需要优先复核的经营对象数量"},
    ])
    add_dataframe_to_docx(doc, summary_rows, max_rows=10)

    add_heading(doc, "三、经营表现与结构洞察", 1)
    if date_col:
        trend = build_trend(df, date_col, main_metric)
        add_heading(doc, "3.1 时间趋势分析", 2)
        add_para(doc, trend_interpretation(trend, main_metric), indent=True)
        if len(trend) >= 2:
            add_mpl_figure_to_docx(doc, make_mpl_trend_figure(trend, main_metric), f"图1 {main_metric}时间趋势")
            add_para(doc, "图表解读：趋势图用于判断主指标在不同周期的方向变化、阶段性高低点以及近期是否存在持续上升或回落。", indent=True)
    else:
        add_heading(doc, "3.1 结构分布分析", 2)
        add_para(doc, f"当前数据不满足时间趋势分析条件，因此重点观察{main_metric}在不同记录中的集中程度、离散程度和高低值分布。", indent=True)
        add_mpl_figure_to_docx(doc, make_mpl_distribution_figure(df, main_metric), f"图1 {main_metric}分布情况")
        _add_distribution_interpretation(doc, df, main_metric)

    if dimensions:
        add_heading(doc, "3.2 多维结构贡献", 2)
        add_para(doc, f"多维结构分析用于回答“{main_metric}主要由哪些业务维度贡献”。系统优先选择“{dimensions[0]}”进行代表性展示，并可在系统页面继续下钻其他维度。", indent=True)
        add_mpl_figure_to_docx(doc, make_mpl_dimension_bar_figure(df, dimensions[0], main_metric), f"图2 按{dimensions[0]}汇总{main_metric}")
        _add_dimension_interpretation(doc, df, dimensions, main_metric)
    else:
        add_para(doc, "当前数据中可用于分组下钻的维度字段较少，建议补充地区、部门、客户类型、产品类别等字段，以提升结构分析效果。", indent=True)

    add_heading(doc, "四、重点风险与异常核查", 1)
    if anomaly_df is not None and len(anomaly_df) and '是否经营异常' in anomaly_df.columns:
        risk = _risk_profile(anomaly_df)
        add_para(doc, f"系统基于主指标偏离、多指标组合偏离、业务规则风险和模型异常贡献识别异常经营单元。当前共识别出{risk['abnormal_n']}个异常经营单元，其中高风险单元{risk['high_n']}个。该结果用于确定核查优先级，而不是直接给出最终业务定性。", indent=True)
        add_mpl_figure_to_docx(doc, make_mpl_risk_figure(anomaly_df), "图3 风险等级分布")
        _add_risk_interpretation(doc, anomaly_df, main_metric)
        _add_risk_review_section(doc, anomaly_df, main_metric, dimensions, section_title="4.1 重点异常核查清单")
    else:
        add_para(doc, "当前数据不足以进行稳定的异常诊断，建议补充更多数值指标和业务维度后再进行风险识别。", indent=True)

    add_heading(doc, "五、管理建议与下一步行动", 1)
    action_df = _build_management_action_table(main_metric, dimensions, anomaly_df)
    add_dataframe_to_docx(doc, action_df, max_rows=10)
    add_para(doc, "以上建议用于辅助管理层确定核查优先级。实际决策仍需结合企业业务背景、政策制度、预算目标和原始业务单据进行综合判断。", indent=True)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================
# 页面入口、功能导航与数据初始化
# ============================================================

st.markdown("""
<div class="hero">
    <h1>智策经营——AI 驱动的多维经营分析与决策支持系统</h1>
    <p>先理解字段，再自适应分析：系统基于上传数据自动完成字段解释、经营态势、经营洞察、异常诊断、问数助手与管理层简报生成。</p>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("## 🧭 功能导航")
st.sidebar.markdown("""
- 🧩 **数据治理**：字段识别、质量评估、字段资产地图  
- 🏠 **经营态势**：核心指标、趋势演示、风险摘要  
- 📊 **经营洞察**：维度对比、贡献度、指标关系、热力图  
- ⚠️ **风险识别**：异常检测、风险评分、AI解释  
- 💬 **问数助手**：自然语言问数、SQL校验、模型对比  
- 📝 **决策简报**：自动生成管理层经营分析 Word 报告  
""")

st.sidebar.markdown("---")
st.sidebar.markdown("## 📂 数据上传")
st.sidebar.markdown("支持 Excel / CSV，系统会根据上传数据动态识别字段和分析口径。")
uploaded_file = st.sidebar.file_uploader(
    "上传经营数据文件",
    type=["xlsx", "xls", "csv"]
)

if uploaded_file is None:
    st.markdown("### 使用流程")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""
        <div class="section-card">
            <h3>① 上传数据</h3>
            <p class="small-text">上传 Excel 或 CSV 经营数据，系统自动识别金额、数量、时间、维度、ID、比例等字段角色。</p>
        </div>
        """, unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div class="section-card">
            <h3>② 智能分析</h3>
            <p class="small-text">围绕主指标生成经营态势、多维洞察、异常风险识别，并解释“为什么异常、哪里值得关注”。</p>
        </div>
        """, unsafe_allow_html=True)
    with c3:
        st.markdown("""
        <div class="section-card">
            <h3>③ 问数与简报</h3>
            <p class="small-text">支持自然语言问数、SQL执行校验、模型对比，并自动生成管理层经营分析 Word 简报。</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("""
    <div class="tip-card">
    <b>开始使用：</b>请在左侧上传经营数据文件。上传后，系统会自动进入数据治理、经营态势、经营洞察、风险识别、问数助手和决策简报模块。
    </div>
    """, unsafe_allow_html=True)
    st.stop()

try:
    raw_df = read_uploaded_file(uploaded_file)
except Exception as e:
    st.error(f"文件读取失败：{e}")
    st.stop()

if raw_df is None or raw_df.empty:
    st.error("上传文件为空，无法开展经营分析。")
    st.stop()

# 标准化字段名，避免空格、特殊符号导致 SQL 或图表报错
raw_df = raw_df.copy()
normalized_cols = []
used_cols = set()
for idx, col in enumerate(raw_df.columns):
    new_col = normalize_column_name(col)
    if not new_col:
        new_col = f"字段{idx + 1}"
    base = new_col
    k = 1
    while new_col in used_cols:
        k += 1
        new_col = f"{base}_{k}"
    used_cols.add(new_col)
    normalized_cols.append(new_col)
raw_df.columns = normalized_cols

# 字段语义识别
meta = infer_field_metadata(raw_df)
metric_candidates = get_metric_candidates(meta)
dimension_candidates = get_dimension_candidates(meta)
date_candidates = get_date_candidates(meta)

# 兜底：如果规则未识别到主指标，则从全表中寻找可解析数值字段
if not metric_candidates:
    for c in raw_df.columns:
        if is_id_like(c):
            continue
        s = parse_numeric_series(raw_df[c], aggressive=True)
        if s.notna().sum() > 0:
            metric_candidates.append(c)

if not metric_candidates:
    st.error("当前数据中未识别到可用于分析的数值指标。请检查文件中是否包含金额、数量、成本、利润、收入等数值字段。")
    st.stop()

st.sidebar.markdown("---")
st.sidebar.markdown("## ⚙️ 字段配置")

main_metric = st.sidebar.selectbox("主分析指标", options=metric_candidates, index=0)
date_choice = st.sidebar.selectbox("日期字段", options=["无"] + date_candidates, index=0)
date_col = None if date_choice == "无" else date_choice

default_nums = list(dict.fromkeys([main_metric] + [c for c in metric_candidates if c != main_metric][:4]))
selected_numeric_cols = st.sidebar.multiselect(
    "参与分析的数值字段",
    options=metric_candidates,
    default=default_nums
)
selected_dimensions = st.sidebar.multiselect(
    "分析维度字段",
    options=dimension_candidates,
    default=dimension_candidates[:3]
)
missing_strategy = st.sidebar.selectbox(
    "缺失值处理策略",
    options=["保留并在分析时忽略", "删除主指标缺失行", "数值中位数填充，类别填未知"],
    index=0
)

if main_metric not in selected_numeric_cols:
    selected_numeric_cols = [main_metric] + selected_numeric_cols

# 所有页面渲染前必须完成初始化，否则 clean_summary / numeric_report / anomaly_df 等变量会未定义
df, quality_report, clean_summary, numeric_report = clean_data(
    raw_df,
    selected_numeric_cols,
    date_col,
    missing_strategy
)

valid_date_col = None
if date_col and date_col in df.columns:
    parsed, ratio, valid, reason = try_parse_date(df[date_col])
    if valid:
        valid_date_col = date_col
        df[date_col] = parsed
    else:
        st.sidebar.warning(f"日期字段“{date_col}”暂不作为有效日期使用：{reason}")

numeric_cols = safe_numeric_cols(df, selected_numeric_cols)

if main_metric not in numeric_cols:
    st.error(f"主指标“{main_metric}”没有可用数值。请在左侧重新选择主指标，或在“数据治理”中查看数值转换报告。")
    st.stop()

selected_dimensions = [d for d in selected_dimensions if d in df.columns]

try:
    anomaly_df = anomaly_detection(df, main_metric, numeric_cols, selected_dimensions, valid_date_col)
except Exception as e:
    st.warning(f"异常诊断暂未生成：{e}")
    anomaly_df = pd.DataFrame()


# ============================================================
# 数据治理可视化辅助函数
# ============================================================

def compute_data_health_score(clean_summary, meta, numeric_report, metric_candidates, dimension_candidates, date_candidates):
    rows = max(clean_summary.get("清洗后行数", 0), 1)
    cols = max(clean_summary.get("清洗后字段数", 0), 1)
    missing = clean_summary.get("缺失单元格数", 0)
    dup = clean_summary.get("重复行数", 0)

    missing_rate = min(missing / max(rows * cols, 1), 1)
    dup_rate = min(dup / max(clean_summary.get("原始行数", rows), 1), 1)

    missing_score = max(0, 100 - missing_rate * 220)
    dup_score = max(0, 100 - dup_rate * 260)

    if numeric_report is not None and len(numeric_report) > 0 and "转换后有效数值数" in numeric_report.columns and "转换前非空数" in numeric_report.columns:
        valid_rates = []
        for _, r in numeric_report.iterrows():
            before = float(r.get("转换前非空数", 0))
            after = float(r.get("转换后有效数值数", 0))
            valid_rates.append(after / before if before > 0 else 1)
        numeric_score = np.mean(valid_rates) * 100 if valid_rates else 80
    else:
        numeric_score = 80

    structure_score = 100 if metric_candidates and dimension_candidates else (72 if metric_candidates else 45)
    date_score = 100 if date_candidates else 82

    score = 0.28 * missing_score + 0.18 * dup_score + 0.24 * numeric_score + 0.20 * structure_score + 0.10 * date_score
    return int(round(max(0, min(score, 100))))


def health_level(score):
    if score >= 88:
        return "优秀", "数据结构较完整，适合直接开展经营分析。"
    if score >= 75:
        return "良好", "数据整体可用，建议关注少量质量问题。"
    if score >= 60:
        return "可用", "数据可以分析，但建议先处理缺失、重复或字段口径问题。"
    return "待治理", "数据质量或字段结构存在明显不足，建议先完成数据治理。"


def gauge_chart(score):
    level, desc = health_level(score)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        number={"suffix": "分", "font": {"size": 42}},
        title={"text": f"数据健康评分｜{level}", "font": {"size": 22}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 1},
            "bar": {"thickness": 0.28},
            "steps": [
                {"range": [0, 60], "color": "#FCE8E6"},
                {"range": [60, 75], "color": "#FFF4D6"},
                {"range": [75, 88], "color": "#E8F3FF"},
                {"range": [88, 100], "color": "#E7F7EF"},
            ],
            "threshold": {"line": {"width": 4}, "thickness": 0.75, "value": score}
        }
    ))
    fig.update_layout(height=300, margin=dict(l=20, r=20, t=55, b=20), paper_bgcolor="#FFFFFF", font=dict(family="Microsoft YaHei"))
    st.plotly_chart(fig, use_container_width=True)
    return level, desc





def management_status_summary(df, main_metric, anomaly_df, trend=None, selected_dimensions=None):
    """经营态势页摘要卡：让用户先看到当前经营状态，而不是只看数字。"""
    metric_series = pd.to_numeric(df[main_metric], errors="coerce").dropna()

    if trend is not None and len(trend) >= 2 and main_metric in trend.columns:
        first = pd.to_numeric(trend[main_metric], errors="coerce").iloc[0]
        last = pd.to_numeric(trend[main_metric], errors="coerce").iloc[-1]
        change = (last - first) / abs(first) if first != 0 else np.nan
        if pd.isna(change):
            status = "趋势待确认"
            status_desc = "当前时间序列不足以形成稳定趋势判断。"
        elif change > 0.1:
            status = "增长态势"
            status_desc = f"末期较初期增长约 {change:.1%}，建议进一步查看增长来源。"
        elif change < -0.1:
            status = "回落态势"
            status_desc = f"末期较初期下降约 {abs(change):.1%}，建议关注拖累维度。"
        else:
            status = "相对平稳"
            status_desc = "期初与期末差异不大，可重点观察结构和异常变化。"
    else:
        status = "结构分析为主"
        status_desc = "当前未形成有效时间趋势，系统将优先从维度贡献、结构集中和异常对象进行分析。"

    if len(metric_series) > 0:
        max_mean = metric_series.max() / metric_series.mean() if metric_series.mean() != 0 else np.nan
        if not pd.isna(max_mean) and max_mean >= 6:
            structure = "长尾集中"
            struct_desc = "少数高值单元对整体结果影响较大。"
        elif not pd.isna(max_mean) and max_mean >= 3:
            structure = "适度集中"
            struct_desc = "存在一定高值集中现象。"
        else:
            structure = "相对均衡"
            struct_desc = "极端高值对整体影响相对有限。"
    else:
        structure = "待观察"
        struct_desc = "主指标有效值不足，暂不形成结构判断。"

    if anomaly_df is not None and len(anomaly_df) and "是否经营异常" in anomaly_df.columns:
        abnormal_rate = anomaly_df["是否经营异常"].sum() / max(len(anomaly_df), 1)
        if abnormal_rate >= 0.1:
            risk = "重点关注"
        elif abnormal_rate > 0:
            risk = "局部关注"
        else:
            risk = "正常"
        risk_desc = f"异常占比约 {abnormal_rate:.1%}。"
    else:
        risk = "待识别"
        risk_desc = "暂未形成异常识别结果。"

    st.markdown(f"""
    <div class="section-card">
        <h3 style="margin-top:0;">管理层态势摘要</h3>
        <div class="asset-grid-compact">
            <div class="asset-card compact">
                <div class="asset-label">当前表现</div>
                <div class="asset-value" style="font-size:26px;">{status}</div>
                <div class="asset-note">{status_desc}</div>
            </div>
            <div class="asset-card compact">
                <div class="asset-label">结构特征</div>
                <div class="asset-value" style="font-size:26px;">{structure}</div>
                <div class="asset-note">{struct_desc}</div>
            </div>
            <div class="asset-card compact">
                <div class="asset-label">风险状态</div>
                <div class="asset-value" style="font-size:26px;">{risk}</div>
                <div class="asset-note">{risk_desc}</div>
            </div>
            <div class="asset-card compact">
                <div class="asset-label">管理提示</div>
                <div class="asset-value" style="font-size:22px;">先看结构，再查异常</div>
                <div class="asset-note">先定位贡献集中的维度，再进入风险识别页复核高风险对象。</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def dimension_treemap_chart(df, dim, metric, topn=20, second_dim=None):
    """
    结构图智能切换：
    - 维度项较少时，用横向排名条形图，更清晰、更适合经营汇报；
    - 维度项数量 >= 8，或存在二级维度时，用 Treemap 展示结构集中和层级贡献。
    """
    if not dim or dim not in df.columns or metric not in df.columns:
        return

    if second_dim == dim or (second_dim is not None and second_dim not in df.columns):
        second_dim = None

    method = aggregation_method_for_metric(metric) if "aggregation_method_for_metric" in globals() else "sum"
    agg_func = "mean" if method == "mean" else "sum"

    if second_dim:
        group_cols = [dim, second_dim]
    else:
        group_cols = [dim]

    g = df.groupby(group_cols, dropna=False)[metric].agg(agg_func).reset_index()
    value_col = metric
    g[value_col] = pd.to_numeric(g[value_col], errors="coerce").fillna(0)
    g = g[g[value_col].abs() > 0]
    if len(g) == 0:
        show_no_chart_reason(f"{dim}结构分析", "该维度下没有可用于结构分析的有效数值。")
        return

    dim_count = int(df[dim].nunique(dropna=False))
    use_treemap = bool(second_dim) or dim_count >= 8

    if use_treemap:
        g = g.sort_values(value_col, key=lambda s: s.abs(), ascending=False).head(topn).copy()
        if second_dim:
            labels = []
            parents = []
            values = []
            # parent nodes
            parent_df = g.groupby(dim, dropna=False)[value_col].sum().reset_index()
            for _, r in parent_df.iterrows():
                labels.append(str(r[dim]))
                parents.append("")
                values.append(abs(float(r[value_col])))
            # child nodes
            for _, r in g.iterrows():
                labels.append(f"{r[second_dim]}")
                parents.append(str(r[dim]))
                values.append(abs(float(r[value_col])))
            title = f"{dim} × {second_dim}结构地图：{metric}贡献分布"
        else:
            labels = g[dim].astype(str).tolist()
            parents = [""] * len(labels)
            values = g[value_col].abs().astype(float).tolist()
            title = f"{dim}结构地图：{metric}贡献分布"

        fig = go.Figure(go.Treemap(
            labels=labels,
            parents=parents,
            values=values,
            textinfo="label+percent root",
            hovertemplate="类别=%{label}<br>贡献值=%{value:,.2f}<br>占比=%{percentRoot:.1%}<extra></extra>",
            marker=dict(line=dict(width=1, color="white"))
        ))
        fig.update_layout(
            title=title,
            margin=dict(l=8, r=8, t=54, b=8)
        )
        st.plotly_chart(chart_layout(fig, 520), use_container_width=True)
        st.caption("结构地图用于观察主指标是否集中在少数维度项上；当维度项较多或存在层级维度时，面积越大表示贡献越高。")
    else:
        plot = g.sort_values(value_col, key=lambda s: s.abs(), ascending=True).tail(topn).copy()
        total = plot[value_col].sum()
        plot["贡献占比"] = plot[value_col] / total if total != 0 else 0
        fig = go.Figure(go.Bar(
            y=plot[dim].astype(str),
            x=plot[value_col],
            orientation="h",
            text=[f"{money_fmt(v)}｜{p:.1%}" for v, p in zip(plot[value_col], plot["贡献占比"])],
            textposition="outside",
            hovertemplate=f"{dim}=%{{y}}<br>{metric}=%{{x:,.2f}}<extra></extra>",
            marker=dict(color="#1E77D3", opacity=0.82)
        ))
        fig.update_layout(
            title=f"{dim}贡献排名：{metric}结构分布",
            xaxis_title=metric,
            yaxis_title=dim,
            margin=dict(l=20, r=80, t=58, b=30)
        )
        st.plotly_chart(chart_layout(fig, 500), use_container_width=True)

        top_row = plot.sort_values(value_col, key=lambda s: s.abs(), ascending=False).head(1)
        if len(top_row):
            top_name = str(top_row.iloc[0][dim])
            top_share = float(top_row.iloc[0]["贡献占比"])
            st.caption(f"当前维度项较少，系统自动采用贡献排名图。{top_name}贡献最高，占比约 {top_share:.1%}，可进一步下钻其客户结构、利润表现和异常风险。")

def contribution_waterfall_chart(comp, dim, metric, p0, p1):
    if comp is None or len(comp) == 0 or "变化额" not in comp.columns:
        return
    plot = comp.sort_values("变化额", key=lambda s: s.abs(), ascending=False).head(12).copy()
    total_change = plot["变化额"].sum()
    fig = go.Figure(go.Waterfall(
        name="变化贡献",
        orientation="v",
        measure=["relative"] * len(plot) + ["total"],
        x=plot[dim].astype(str).tolist() + ["合计变化"],
        y=plot["变化额"].astype(float).tolist() + [total_change],
        text=[money_fmt(v) for v in plot["变化额"].astype(float).tolist()] + [money_fmt(total_change)],
        textposition="outside"
    ))
    fig.update_layout(
        title=f"{p0} 至 {p1} 的{metric}变化归因瀑布图",
        xaxis_title=dim,
        yaxis_title="变化额",
        showlegend=False
    )
    st.plotly_chart(chart_layout(fig, 520), use_container_width=True)
    st.caption("瀑布图用于解释主指标变化由哪些维度项拉动或拖累，正值代表拉动增长，负值代表拖累下降。")


def risk_matrix_chart(anomaly_df, main_metric):
    """风险矩阵：用影响程度 × 异常程度识别优先复核对象。"""
    if anomaly_df is None or len(anomaly_df) == 0 or "风险得分" not in anomaly_df.columns or main_metric not in anomaly_df.columns:
        return
    plot = anomaly_df.copy()
    plot[main_metric] = pd.to_numeric(plot[main_metric], errors="coerce")
    plot["风险得分"] = pd.to_numeric(plot["风险得分"], errors="coerce")
    plot = plot.dropna(subset=[main_metric, "风险得分"])
    if len(plot) < 3:
        st.info("当前异常样本不足，暂不生成风险矩阵。")
        return

    plot["影响程度"] = plot[main_metric].abs().rank(pct=True) * 100
    if "风险等级" not in plot.columns:
        plot["风险等级"] = "未分级"

    size_raw = plot[main_metric].abs().replace([np.inf, -np.inf], np.nan).fillna(0)
    if size_raw.max() > size_raw.min():
        sizes = 9 + (size_raw - size_raw.min()) / (size_raw.max() - size_raw.min()) * 28
    else:
        sizes = pd.Series([17] * len(plot), index=plot.index)

    fig = go.Figure()

    # 四象限背景：帮助用户理解坐标含义
    zones = [
        (0, 70, 0, 70, "观察区", "#F4F8FC"),
        (70, 100, 0, 70, "高影响观察区", "#EEF6FF"),
        (0, 70, 70, 100, "异常关注区", "#FFF7E8"),
        (70, 100, 70, 100, "优先复核区", "#FFECEC"),
    ]
    for x0, x1, y0, y1, _, color in zones:
        fig.add_shape(
            type="rect", x0=x0, x1=x1, y0=y0, y1=y1,
            fillcolor=color, opacity=0.65, layer="below", line_width=0
        )

    for level, sub in plot.groupby("风险等级", dropna=False):
        idx = sub.index
        hover_text = []
        for _, r in sub.iterrows():
            parts = [
                f"风险等级={r.get('风险等级', '')}",
                f"风险得分={r.get('风险得分', np.nan):.1f}",
                f"影响程度={r.get('影响程度', np.nan):.1f}",
                f"{main_metric}={money_fmt(r.get(main_metric, np.nan))}"
            ]
            if "异常依据" in r:
                parts.append(f"异常依据={str(r.get('异常依据', ''))[:100]}")
            hover_text.append("<br>".join(parts))
        fig.add_trace(go.Scatter(
            x=sub["影响程度"],
            y=sub["风险得分"],
            mode="markers",
            name=str(level),
            marker=dict(size=sizes.loc[idx], opacity=0.78, line=dict(width=1, color="white")),
            text=hover_text,
            hovertemplate="%{text}<extra></extra>"
        ))

    fig.add_hline(y=70, line_dash="dash", line_color="#D9822B", line_width=2)
    fig.add_vline(x=70, line_dash="dash", line_color="#D9822B", line_width=2)

    annotations = [
        dict(x=35, y=35, text="观察区", showarrow=False, font=dict(size=13, color="#74839A")),
        dict(x=85, y=35, text="高影响<br>观察区", showarrow=False, font=dict(size=13, color="#1E77D3")),
        dict(x=35, y=86, text="异常<br>关注区", showarrow=False, font=dict(size=13, color="#A76700")),
        dict(x=85, y=86, text="优先复核区", showarrow=False, font=dict(size=15, color="#B42318")),
    ]

    fig.update_layout(
        title="风险矩阵：影响程度 × 异常程度",
        xaxis_title="影响程度（主指标分位，越靠右影响越大）",
        yaxis_title="异常程度（风险得分，越靠上异常越强）",
        xaxis=dict(range=[0, 102], zeroline=False),
        yaxis=dict(range=[0, 102], zeroline=False),
        annotations=annotations,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
    )
    st.plotly_chart(chart_layout(fig, 540), use_container_width=True)
    st.caption("风险矩阵用于确定核查优先级：右上角代表同时具有较高经营影响和较高异常程度的对象，应优先复核；左下角通常为低优先级观察对象。")

def risk_source_summary_cards(anomaly_df):
    if anomaly_df is None or len(anomaly_df) == 0:
        return
    score_cols = ["主指标偏离", "多指标组合偏离", "业务规则风险", "模型异常贡献"]
    valid = [c for c in score_cols if c in anomaly_df.columns]
    if not valid:
        return

    mean_scores = anomaly_df[valid].apply(pd.to_numeric, errors="coerce").mean().sort_values(ascending=False)
    top_source = mean_scores.index[0] if len(mean_scores) else "暂无"

    high_count = 0
    if "风险等级" in anomaly_df.columns:
        high_count = int((anomaly_df["风险等级"].astype(str).str.contains("高", na=False)).sum())
    total = len(anomaly_df)

    desc_map = {
        "主指标偏离": "当前风险主要来自主指标处于异常高位或低位，建议核查高值/低值经营单元是否符合真实业务表现。",
        "多指标组合偏离": "当前风险主要来自多个指标同步偏离，建议结合销售额、数量、利润、折扣等字段交叉复核。",
        "业务规则风险": "当前风险主要来自利润、折扣、成本或比例类业务规则信号，建议核查政策口径和核算逻辑。",
        "模型异常贡献": "当前风险主要来自多指标组合的离群程度，建议优先查看风险矩阵右上角对象。"
    }

    st.markdown(f"""
    <div class="section-card">
        <h3 style="margin-top:0;">风险画像摘要</h3>
        <p class="small-text"><b>高风险对象：</b>{high_count} / {total} 个。</p>
        <p class="small-text"><b>主要风险来源：</b>{top_source}。{desc_map.get(top_source, "需要结合明细进一步复核。")}</p>
        <p class="small-text"><b>使用建议：</b>先根据风险矩阵定位右上角对象，再结合异常依据、原始单据、折扣政策、利润口径和客户/渠道背景进行复核。</p>
    </div>
    """, unsafe_allow_html=True)

def query_process_visual():
    steps = [
        ("用户问题", "自然语言输入"),
        ("字段解析", "匹配指标/维度/时间"),
        ("SQL生成", "调用大模型生成查询"),
        ("数据库执行", "在上传数据上执行"),
        ("系统校验", "检查字段/结果/评分"),
        ("AI回答", "输出结论与建议"),
    ]
    html = '<div class="section-card"><h3 style="margin-top:0;">问数过程可信链路</h3><div style="display:flex;gap:10px;flex-wrap:wrap;">'
    for i, (title, desc) in enumerate(steps, 1):
        html += f'<div style="flex:1;min-width:135px;background:#F3F8FF;border:1px solid #D9EAFF;border-radius:16px;padding:12px 14px;"><div style="font-size:13px;color:#1E77D3;font-weight:900;">STEP {i}</div><div style="font-size:16px;color:#10213F;font-weight:950;margin-top:4px;">{title}</div><div style="font-size:13px;color:#66758C;margin-top:5px;">{desc}</div></div>'
    html += '</div></div>'
    st.markdown(html, unsafe_allow_html=True)


def recommend_followup_questions(question, main_metric, selected_dimensions, numeric_cols):
    dim = selected_dimensions[0] if selected_dimensions else "主要维度"
    qlist = [
        f"{main_metric}最高的{dim}是什么？",
        f"哪些{dim}的{main_metric}增长最快？",
        f"{main_metric}较高的对象是否存在异常风险？",
    ]
    # 加入利润/折扣类追问
    for c in numeric_cols:
        if "利润" in str(c) or "收益" in str(c):
            qlist.append(f"{main_metric}高但{c}低的对象有哪些？")
            break
    for c in numeric_cols:
        if is_rate_like(c) or "折扣" in str(c):
            qlist.append(f"{c}较高的对象对{main_metric}有什么影响？")
            break
    return qlist[:5]


def governance_health_gauge(health_score):
    level, desc = health_level(health_score)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=float(health_score),
        number={"suffix": " 分", "font": {"size": 30, "color": "#071F43"}},
        title={"text": "数据健康度", "font": {"size": 18, "color": "#10213F"}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#DDE6F1", "tickfont": {"size": 12}},
            "bar": {"color": "#1E77D3", "thickness": 0.24},
            "bgcolor": "white",
            "borderwidth": 1,
            "bordercolor": "#E5EBF3",
            "steps": [
                {"range": [0, 60], "color": "#FFE8E8"},
                {"range": [60, 80], "color": "#FFF4D8"},
                {"range": [80, 100], "color": "#EAF8F0"}
            ],
            "threshold": {
                "line": {"color": "#123A63", "width": 3},
                "thickness": 0.75,
                "value": float(health_score)
            }
        },
        domain={"x": [0, 1], "y": [0.05, 0.9]}
    ))
    fig.update_layout(
        height=250,
        margin=dict(l=10, r=10, t=42, b=8),
        paper_bgcolor="#FFFFFF",
        font=dict(family="Microsoft YaHei, PingFang SC, Arial", size=14)
    )
    st.plotly_chart(fig, use_container_width=True)
    st.markdown(
        f'<div class="detail-expander-note" style="text-align:center;margin-top:-12px;"><b>{level}</b>｜{desc}</div>',
        unsafe_allow_html=True
    )

def governance_asset_overview_cards(health_score, clean_summary, meta, metric_candidates, dimension_candidates, date_candidates, numeric_report, quality_report):
    rows = int(clean_summary.get("清洗后行数", 0))
    cols = int(clean_summary.get("清洗后字段数", 0))
    valid_metric = len(metric_candidates)
    valid_dim = len(dimension_candidates)
    valid_date = len(date_candidates)

    issue_total = int(clean_summary.get("重复行数", 0)) + int(clean_summary.get("缺失单元格数", 0))
    if numeric_report is not None and len(numeric_report) and "缺失/无法解析数" in numeric_report.columns:
        issue_total += int(pd.to_numeric(numeric_report["缺失/无法解析数"], errors="coerce").fillna(0).sum())

    html = (
        '<div class="asset-grid-compact">'
        f'<div class="asset-card compact"><div class="asset-label">可分析记录</div><div class="asset-value">{rows:,}</div><div class="asset-note">清洗后进入分析的数据行数</div></div>'
        f'<div class="asset-card compact"><div class="asset-label">字段资产</div><div class="asset-value">{cols}</div><div class="asset-note">系统已识别的字段总数</div></div>'
        f'<div class="asset-card compact"><div class="asset-label">分析路径</div><div class="asset-value">{valid_metric}/{valid_dim}/{valid_date}</div><div class="asset-note">主指标 / 维度 / 日期字段</div></div>'
        f'<div class="asset-card compact"><div class="asset-label">质量关注点</div><div class="asset-value">{issue_total:,}</div><div class="asset-note">缺失、重复或解析问题线索</div></div>'
        '</div>'
    )
    st.markdown(html, unsafe_allow_html=True)

def governance_flow_summary(health_score, main_metric, metric_candidates, dimension_candidates, date_candidates, clean_summary):
    level, desc = health_level(health_score)
    metric_text = "、".join(metric_candidates[:4]) if metric_candidates else "暂无明确主指标"
    dim_text = "、".join(dimension_candidates[:4]) if dimension_candidates else "暂无明确维度字段"
    date_text = "、".join(date_candidates[:2]) if date_candidates else "未识别到有效日期字段"
    rows = clean_summary.get("清洗后行数", 0)
    cols = clean_summary.get("清洗后字段数", 0)
    st.markdown(f"""
    <div class="governance-hero">
        <div class="governance-hero-title">数据资产总览</div>
        <div class="governance-hero-text">
            系统已将上传文件转换为可分析数据资产：当前共有 <b>{rows:,}</b> 条可分析记录、<b>{cols}</b> 个字段，
            数据健康状态为 <b>{level}</b>。推荐以 <b>{main_metric}</b> 作为当前主分析指标。
        </div>
        <div style="margin-top:12px;">
            <span class="governance-step">主指标：{metric_text}</span>
            <span class="governance-step">分析维度：{dim_text}</span>
            <span class="governance-step">时间字段：{date_text}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)


def governance_action_center(clean_summary, numeric_report, quality_report, metric_candidates, dimension_candidates, date_candidates):
    actions = []

    missing = int(clean_summary.get("缺失单元格数", 0))
    duplicate = int(clean_summary.get("重复行数", 0))

    numeric_bad = 0
    if numeric_report is not None and len(numeric_report) and "缺失/无法解析数" in numeric_report.columns:
        numeric_bad = int(pd.to_numeric(numeric_report["缺失/无法解析数"], errors="coerce").fillna(0).sum())

    if missing > 0:
        actions.append(("缺失值治理", f"识别到 {missing:,} 个缺失单元格，建议结合业务口径选择忽略、填充或人工复核。", "需关注"))
    else:
        actions.append(("缺失值治理", "当前未发现明显缺失单元格，字段完整性较好。", "正常"))

    if duplicate > 0:
        actions.append(("重复记录治理", f"识别到 {duplicate:,} 行重复记录，建议核查是否为重复导入或真实重复业务。", "需关注"))
    else:
        actions.append(("重复记录治理", "当前未发现重复行，记录唯一性较好。", "正常"))

    if numeric_bad > 0:
        actions.append(("数值格式治理", f"存在 {numeric_bad:,} 个数值缺失或无法解析单元，可能影响图表、问数和风险识别。", "需关注"))
    else:
        actions.append(("数值格式治理", "数值字段可解析性较好，适合开展指标分析。", "正常"))

    if not date_candidates:
        actions.append(("时间分析能力", "暂未识别到有效日期字段，系统将优先开展结构、分布和多维对比分析。", "提示"))
    else:
        actions.append(("时间分析能力", f"识别到 {len(date_candidates)} 个日期字段，可支持趋势、环比和阶段变化分析。", "正常"))

    if not metric_candidates or not dimension_candidates:
        actions.append(("分析路径完整性", "主指标或维度字段不足，建议补充金额、数量、客户、地区、类别等字段。", "需关注"))
    else:
        actions.append(("分析路径完整性", "主指标与维度字段较完整，可支持经营态势、多维洞察和风险识别。", "正常"))

    html_parts = ['<div class="action-list-card"><h3 style="margin-top:0;">质量体检与处理建议</h3>']
    for title, desc, status in actions:
        cls = "warn" if status == "需关注" else ("ok" if status == "正常" else "")
        html_parts.append(
            f'<div class="action-item">'
            f'<div><div class="action-title">{title}</div><div class="action-desc">{desc}</div></div>'
            f'<div class="action-badge {cls}">{status}</div>'
            f'</div>'
        )
    html_parts.append("</div>")
    st.markdown("".join(html_parts), unsafe_allow_html=True)

def field_asset_table(meta, metric_candidates, dimension_candidates, date_candidates):
    show_cols = ["字段名", "字段含义解释", "字段类型", "推荐角色", "推荐聚合方式", "是否适合作为主指标", "注意事项"]
    cols = [c for c in show_cols if c in meta.columns]
    if not cols:
        return meta
    out = meta[cols].copy()
    role_order = {"度量指标": 0, "数值字段": 1, "维度字段": 2, "时间字段": 3, "标识字段": 4, "文本字段": 5}
    out["_排序"] = out["推荐角色"].map(role_order).fillna(99) if "推荐角色" in out.columns else 99
    return out.sort_values("_排序").drop(columns=["_排序"], errors="ignore")




def role_distribution_chart(meta):
    role_count = meta["推荐角色"].value_counts().reset_index()
    role_count.columns = ["字段角色", "字段数量"]
    order = ["度量指标", "数值字段", "维度字段", "时间字段", "进度/时长字段", "标识字段", "文本字段"]
    role_count["排序"] = role_count["字段角色"].apply(lambda x: order.index(x) if x in order else 99)
    role_count = role_count.sort_values(["排序", "字段数量"], ascending=[True, False])

    fig = go.Figure(go.Pie(
        labels=role_count["字段角色"],
        values=role_count["字段数量"],
        hole=0.56,
        textinfo="label+percent",
        hovertemplate="%{label}<br>字段数量=%{value}<br>占比=%{percent}<extra></extra>"
    ))
    fig.update_layout(
        title="字段角色资产分布",
        annotations=[dict(text=f"{int(role_count['字段数量'].sum())}<br>字段", x=0.5, y=0.5, showarrow=False, font=dict(size=22, color="#10213F"))],
        showlegend=True
    )
    st.plotly_chart(chart_layout(fig, 500), use_container_width=True)

def data_quality_issue_chart(clean_summary, quality_report, numeric_report):
    duplicate = int(clean_summary.get("重复行数", 0))
    missing = int(clean_summary.get("缺失单元格数", 0))

    date_fail = 0
    if quality_report is not None and len(quality_report):
        q = quality_report[quality_report["处理环节"].astype(str).str.contains("日期字段解析", na=False)]
        if len(q):
            date_fail = int(pd.to_numeric(q["处理数量"], errors="coerce").fillna(0).sum())

    numeric_fail = 0
    if numeric_report is not None and len(numeric_report) and "缺失/无法解析数" in numeric_report.columns:
        numeric_fail = int(pd.to_numeric(numeric_report["缺失/无法解析数"], errors="coerce").fillna(0).sum())

    issues = pd.DataFrame({
        "问题类型": ["缺失值", "数值无法解析", "重复记录", "日期解析失败"],
        "数量": [missing, numeric_fail, duplicate, date_fail]
    }).sort_values("数量", ascending=True)

    fig = go.Figure(go.Bar(
        x=issues["数量"],
        y=issues["问题类型"],
        orientation="h",
        text=issues["数量"].map(lambda x: f"{int(x):,}"),
        textposition="outside",
        hovertemplate="%{y}<br>数量=%{x:,}<extra></extra>"
    ))
    fig.update_layout(
        title="数据质量问题分布",
        xaxis_title="问题数量",
        yaxis_title="问题类型"
    )
    st.plotly_chart(chart_layout(fig, 500), use_container_width=True)

def make_pills(items, css_class=""):
    if not items:
        return '<span class="field-pill muted">暂无</span>'
    pills = []
    for item in items[:12]:
        pills.append(f'<span class="field-pill {css_class}">{item}</span>')
    if len(items) > 12:
        pills.append(f'<span class="field-pill muted">+{len(items)-12}</span>')
    return '<div class="pill-wrap">' + "".join(pills) + "</div>"


def field_recommendation_cards(meta, metric_candidates, dimension_candidates, date_candidates):
    id_fields = meta[meta["推荐角色"].isin(["标识字段"])]["字段名"].tolist()
    cautious_fields = meta[(meta["字段类型"].astype(str).str.contains("比例|比率", na=False)) | (meta["注意事项"].astype(str).str.contains("谨慎|不能求和|加权", na=False))]["字段名"].tolist()

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""
        <div class="field-card">
            <div class="field-card-title">🎯 推荐主指标</div>
            <div class="field-card-desc">适合作为经营表现、规模、成本、收入或利润的核心分析对象。</div>
            {make_pills(metric_candidates)}
        </div>
        """, unsafe_allow_html=True)
    with c2:
        st.markdown(f"""
        <div class="field-card">
            <div class="field-card-title">🧭 推荐分析维度</div>
            <div class="field-card-desc">适合用于分组对比、结构分析、贡献度分析和下钻定位。</div>
            {make_pills(dimension_candidates)}
        </div>
        """, unsafe_allow_html=True)
    with c3:
        st.markdown(f"""
        <div class="field-card">
            <div class="field-card-title">⏱️ 时间字段</div>
            <div class="field-card-desc">存在时间字段时，可开展趋势、环比、阶段变化和贡献度分析。</div>
            {make_pills(date_candidates)}
        </div>
        """, unsafe_allow_html=True)
    with c4:
        st.markdown(f"""
        <div class="field-card">
            <div class="field-card-title">⚠️ 谨慎使用</div>
            <div class="field-card-desc">系统会区分规模型指标与水平型指标，自动推荐更适合的分析口径。</div>
            {make_pills(list(dict.fromkeys(id_fields + cautious_fields)), "warn")}
        </div>
        """, unsafe_allow_html=True)


def numeric_validity_progress(numeric_report):
    if numeric_report is None or len(numeric_report) == 0:
        st.info("暂无数值转换检查结果。")
        return

    rows = []
    for _, r in numeric_report.iterrows():
        field = r.get("字段", "")
        before = float(r.get("转换前非空数", 0))
        after = float(r.get("转换后有效数值数", 0))
        rate = after / before if before > 0 else 1
        rows.append({"字段": field, "有效率": max(0, min(rate, 1)), "有效率文本": f"{max(0, min(rate, 1)):.1%}"})

    df_rate = pd.DataFrame(rows).sort_values("有效率", ascending=True).tail(8)
    html_parts = ['<div class="compact-progress-wrap">']
    for _, r in df_rate.iterrows():
        pct = float(r["有效率"]) * 100
        html_parts.append(
            f'<div class="compact-progress-row">'
            f'<div class="compact-progress-name">{r["字段"]}</div>'
            f'<div class="compact-progress-track"><div class="compact-progress-fill" style="width:{pct:.1f}%"></div></div>'
            f'<div class="compact-progress-value">{r["有效率文本"]}</div>'
            f'</div>'
        )
    html_parts.append("</div>")
    st.markdown("".join(html_parts), unsafe_allow_html=True)

def chart_data_diagnostic(df, main_metric, numeric_cols, dimensions):
    rows = []
    check_cols = list(dict.fromkeys([main_metric] + list(numeric_cols)))
    for c in check_cols:
        if c not in df.columns:
            continue
        s = to_finite_numeric_series(df, c)
        rows.append({
            "字段": c,
            "有效数值数": int(len(s)),
            "唯一数值数": int(s.nunique(dropna=True)) if len(s) else 0,
            "最小值": float(s.min()) if len(s) else np.nan,
            "最大值": float(s.max()) if len(s) else np.nan,
            "是否可绘图": "是" if len(s) > 0 else "否"
        })
    return pd.DataFrame(rows)


def data_understanding_summary(score, main_metric, metric_candidates, dimension_candidates, date_candidates, clean_summary):
    governance_flow_summary(score, main_metric, metric_candidates, dimension_candidates, date_candidates, clean_summary)

# ============================================================
# 主功能页签
# ============================================================

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🧩 数据治理",
    "🏠 经营态势",
    "📊 经营洞察",
    "⚠️ 风险识别",
    "💬 问数助手",
    "📝 决策简报"
])

# ============================================================
# Tab 1
# ============================================================


with tab1:
    st.subheader("🧩 数据治理")
    st.markdown(
        '<div class="tip-card">系统自动完成数据质量评估、字段角色识别与分析口径推荐，将上传文件转换为可分析、可追溯、可配置的数据资产。</div>',
        unsafe_allow_html=True
    )

    health_score = compute_data_health_score(clean_summary, meta, numeric_report, metric_candidates, dimension_candidates, date_candidates)

    # 顶部：结论摘要 + 健康度仪表盘 + 字段资产卡片
    data_understanding_summary(health_score, main_metric, metric_candidates, dimension_candidates, date_candidates, clean_summary)
    left, right = st.columns([0.88, 1.45])
    with left:
        governance_health_gauge(health_score)
    with right:
        governance_asset_overview_cards(health_score, clean_summary, meta, metric_candidates, dimension_candidates, date_candidates, numeric_report, quality_report)

    st.markdown("### 字段资产地图")
    c1, c2 = st.columns([0.95, 1.35])
    with c1:
        role_distribution_chart(meta)
    with c2:
        field_recommendation_cards(meta, metric_candidates, dimension_candidates, date_candidates)

    st.markdown("### 数据质量体检")
    c1, c2 = st.columns([1.05, 1.15])
    with c1:
        data_quality_issue_chart(clean_summary, quality_report, numeric_report)
    with c2:
        governance_action_center(clean_summary, numeric_report, quality_report, metric_candidates, dimension_candidates, date_candidates)

    st.markdown("### 数值字段可解析性")
    st.markdown('<div class="detail-expander-note">用于判断金额、数量、利润、折扣等字段是否能被系统稳定识别为数值。这里只保留关键字段的有效率概览，完整转换明细可在下方展开查看。</div>', unsafe_allow_html=True)
    numeric_validity_progress(numeric_report)

    st.markdown("### 当前分析口径")
    config_df = pd.DataFrame({
        "配置项": ["主分析指标", "有效日期字段", "分析维度字段", "数值字段", "缺失值策略"],
        "当前选择": [
            main_metric,
            valid_date_col if valid_date_col else "无",
            "、".join(selected_dimensions) if selected_dimensions else "无",
            "、".join(numeric_cols) if numeric_cols else "无",
            missing_strategy
        ],
        "用途说明": [
            "用于经营表现、趋势、贡献度、风险诊断和问数分析",
            "用于时间趋势、阶段变化和动态演示；无有效日期时系统转为结构分析",
            "用于分组对比、贡献度分析、热力图和风险下钻",
            "用于指标关系、异常识别和自然语言问数",
            "用于明确系统对缺失值的处理口径"
        ]
    })
    st.dataframe(config_df, use_container_width=True, hide_index=True)

    st.markdown("### 字段资产清单")
    st.markdown('<div class="detail-expander-note">下表展示字段含义解释、系统推荐角色和聚合口径。系统会区分规模型指标与水平型指标，使后续图表、问数和简报更符合经营管理语境。</div>', unsafe_allow_html=True)
    st.dataframe(field_asset_table(meta, metric_candidates, dimension_candidates, date_candidates).head(20), use_container_width=True, height=420)

    st.markdown("### 明细追溯")
    with st.expander("查看完整字段语义识别表", expanded=False):
        st.dataframe(meta, use_container_width=True, height=420)

    with st.expander("查看数据质量处理日志", expanded=False):
        st.dataframe(quality_report, use_container_width=True, height=360)

    with st.expander("查看数值转换检查", expanded=False):
        st.dataframe(numeric_report, use_container_width=True, height=360)

    with st.expander("查看图表可绘制性诊断", expanded=False):
        st.markdown('<div class="detail-expander-note">如果某张图无法形成有效展示，优先查看这里。有效数值数为0或唯一值过少时，系统会提示原因，而不是显示空白图。</div>', unsafe_allow_html=True)
        st.dataframe(chart_data_diagnostic(df, main_metric, numeric_cols, selected_dimensions), use_container_width=True, height=360)

    with st.expander("查看清洗后数据预览", expanded=False):
        st.dataframe(df.head(30), use_container_width=True, height=360)


# ============================================================
# Tab 2
# ============================================================

with tab2:
    st.subheader("🏠 经营态势")

    abnormal_count = int(anomaly_df["是否经营异常"].sum()) if len(anomaly_df) else 0
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        kpi_card(f"{main_metric}合计", money_fmt(df[main_metric].sum()), "基于当前上传数据")
    with c2:
        kpi_card(f"{main_metric}均值", money_fmt(df[main_metric].mean()), "反映平均水平")
    with c3:
        kpi_card(f"{main_metric}最大值", money_fmt(df[main_metric].max()), "用于识别高值记录")
    with c4:
        kpi_card("异常经营单元", f"{abnormal_count}", "中高风险单元数量")

    overall_trend_for_summary = build_trend(df, valid_date_col, main_metric) if valid_date_col else None
    management_status_summary(df, main_metric, anomaly_df, overall_trend_for_summary, selected_dimensions)

    if valid_date_col:
        st.markdown("### 核心趋势")
        trend_all = build_trend(df, valid_date_col, main_metric)
        periods_all = trend_all["期间"].astype(str).tolist()
        if len(periods_all) >= 2:
            fc1, fc2, fc3 = st.columns([1.2, 1.2, 2])
            with fc1:
                start_p = st.selectbox("趋势起始时间", periods_all, index=0, key="dash_start_period")
            with fc2:
                end_p = st.selectbox("趋势结束时间", periods_all, index=len(periods_all)-1, key="dash_end_period")
            sidx, eidx = periods_all.index(start_p), periods_all.index(end_p)
            if sidx > eidx:
                sidx, eidx = eidx, sidx
            selected_periods = periods_all[sidx:eidx+1]
            trend = trend_all[trend_all["期间"].isin(selected_periods)].copy()
        else:
            selected_periods = periods_all
            trend = trend_all.copy()
        c1, c2 = st.columns([2, 1])
        with c1:
            line_chart(trend, "期间", main_metric, f"{main_metric}时间趋势")
        with c2:
            st.markdown(f'<div class="section-card"><h4>趋势解读</h4><p class="small-text">{trend_interpretation(trend, main_metric)}</p></div>', unsafe_allow_html=True)

        st.markdown("### 动态趋势演示")
        st.markdown("""
        <div class="highlight-note">
        <b>显示条件：</b>只有当上传数据中存在有效日期字段、主指标可按时间汇总，且时间周期不少于 3 个时，系统才会生成动态图。
        动态图按时间顺序逐步展开主指标轨迹，用于观察经营指标从起点到当前周期的演变过程；横轴只保留关键区间刻度，避免时间标签过密。
        </div>
        """, unsafe_allow_html=True)
        animated_trend_chart(trend, main_metric, f"{main_metric}动态趋势演示")

        if selected_dimensions:
            st.markdown("### 分维度趋势对比")
            dc1, dc2, dc3 = st.columns([1.4, 1, 1.2])
            with dc1:
                dim = st.selectbox("选择趋势拆分维度", selected_dimensions, index=0, key="dash_dim_trend")
            with dc2:
                topn_line = st.slider("展示TopN维度项", 3, 12, 6, key="dash_dim_topn")
            temp = periodize(df, valid_date_col)
            if selected_periods:
                temp = temp[temp["_period"].isin(selected_periods)]
            top_dims = df.groupby(dim)[main_metric].sum().sort_values(ascending=False).head(topn_line).index.tolist()
            with dc3:
                highlight = st.selectbox("高亮维度项", ["无"] + [str(x) for x in top_dims], key="dash_highlight_dim")
            dim_trend = temp[temp[dim].isin(top_dims)].groupby(["_period", dim])[main_metric].sum().reset_index()
            fig = go.Figure()
            for name, g in dim_trend.groupby(dim):
                is_high = (highlight != "无" and str(name) == str(highlight))
                fig.add_trace(go.Scatter(x=g["_period"].astype(str), y=g[main_metric], mode="lines+markers", name=str(name), line=dict(width=4 if is_high else 1.6), opacity=1.0 if is_high or highlight == "无" else 0.22, marker=dict(size=7 if is_high else 4)))
            fig.update_layout(title=f"按{dim}对比{main_metric}趋势", xaxis_title="期间", yaxis_title=main_metric, hovermode="closest")
            st.plotly_chart(chart_layout(fig, 460), use_container_width=True)
            st.caption("维度项较多时，系统默认展示主指标贡献较高的TopN。可通过“高亮维度项”突出某一条线，其余线降低透明度，避免画面过乱。")
    else:
        st.markdown('<div class="warn-card">当前未选择有效日期字段，系统已切换为“结构分析模式”：重点展示主指标在不同维度下的分布、排名和异常情况。</div>', unsafe_allow_html=True)

        c1, c2 = st.columns(2)
        with c1:
            histogram_chart(df, main_metric, f"{main_metric}分布")
        with c2:
            box_chart(df, main_metric, f"{main_metric}箱线图")

        if selected_dimensions:
            dim = selected_dimensions[0]
            g = dimension_summary(df, dim, main_metric).head(15)
            c1, c2 = st.columns(2)
            with c1:
                bar_chart(g, dim, f"{main_metric}合计", f"按{dim}对比{main_metric}合计")
            with c2:
                g2 = g.sort_values(f"{main_metric}均值", ascending=False)
                bar_chart(g2, dim, f"{main_metric}均值", f"按{dim}对比{main_metric}均值")

    st.markdown("### 重点风险提示")
    st.markdown("""
    <div class="highlight-note">
    <b>说明：</b>经营态势页只承担“总览预警”作用，帮助管理者快速看到是否存在需要优先关注的高风险对象；
    详细异常原因、风险得分构成和AI解释请进入“风险识别”页面下钻查看。
    </div>
    """, unsafe_allow_html=True)
    if len(anomaly_df):
        top = anomaly_df[anomaly_df["是否经营异常"]].sort_values("风险得分", ascending=False).head(3)
        if len(top):
            ignore_cols = set(numeric_cols + ["记录数", "是否AI异常", "异常得分", "异常依据", "主指标偏离", "多指标组合偏离", "业务规则风险", "模型异常贡献", "风险得分", "风险等级", "是否经营异常"])
            title_cols = [c for c in anomaly_df.columns if c not in ignore_cols]
            rows = []
            for _, row in top.iterrows():
                obj = "｜".join([f"{c}={row[c]}" for c in title_cols[:2]]) if title_cols else f"记录 {row.name}"
                rows.append({
                    "风险对象": obj,
                    "风险等级": row["风险等级"],
                    "风险得分": row["风险得分"],
                    "简要原因": str(row["异常依据"])[:95] + ("..." if len(str(row["异常依据"])) > 95 else "")
                })
            st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
        else:
            st.success("当前未识别出明显中高风险经营单元。")


# ============================================================
# Tab 3
# ============================================================

with tab3:
    st.subheader("📊 经营洞察")
    sub1, sub2, sub3, sub4 = st.tabs(["维度对比", "贡献度分析", "指标关系", "热力图"])

    with sub1:
        if not selected_dimensions:
            st.warning("当前未选择维度字段，无法进行维度对比。")
        else:
            dim = st.selectbox("选择分析维度", selected_dimensions, key="dim_compare")
            g = dimension_summary(df, dim, main_metric)
            c1, c2 = st.columns(2)
            with c1:
                bar_chart(g.head(20), dim, f"{main_metric}合计", f"{dim}维度下{main_metric}合计排名")
            with c2:
                pie_df = g.head(10).copy()
                pie_val_col = f"{main_metric}合计"
                pie_df[pie_val_col] = pd.to_numeric(pie_df[pie_val_col], errors="coerce").replace([np.inf, -np.inf], np.nan)
                pie_df = pie_df.dropna(subset=[pie_val_col])
                if len(pie_df) and pie_df[pie_val_col].abs().sum() > 0:
                    fig = go.Figure(go.Pie(
                        labels=pie_df[dim].astype(str).tolist(),
                        values=pie_df[pie_val_col].astype(float).tolist(),
                        hole=0.35
                    ))
                    fig.update_layout(title=f"{dim}维度贡献占比 Top10")
                    safe_plotly_chart(fig, f"{dim}维度贡献占比 Top10", 460)
                else:
                    show_no_chart_reason(f"{dim}维度贡献占比 Top10", "该维度下没有可用于饼图的正负有效数值。")

            st.markdown("#### 结构地图")
            dimension_treemap_chart(df, dim, main_metric, topn=20, second_dim=next((d for d in selected_dimensions if d != dim), None))

            st.dataframe(g, use_container_width=True)

            importance = dimension_importance(df, selected_dimensions, main_metric)
            if len(importance):
                st.markdown("#### 智能维度推荐")
                st.dataframe(importance, use_container_width=True)
                best_dim = importance.iloc[0]["维度"]
                st.info(f"系统建议优先关注“{best_dim}”维度，因为该维度下主指标的组间差异较明显。")

    with sub2:
        if valid_date_col and selected_dimensions:
            dim = st.selectbox("选择贡献度维度", selected_dimensions, key="contrib_dim")
            periods = sorted(periodize(df, valid_date_col)["_period"].dropna().unique().tolist())
            if len(periods) >= 2:
                pc1, pc2 = st.columns(2)
                with pc1:
                    p0 = st.selectbox("对比起始时间", periods, index=max(0, len(periods)-2), key="contrib_p0")
                with pc2:
                    p1 = st.selectbox("对比结束时间", periods, index=len(periods)-1, key="contrib_p1")
                comp = contribution_analysis(df, valid_date_col, dim, main_metric, p0, p1)
                if len(comp):
                    contribution_waterfall_chart(comp, dim, main_metric, p0, p1)
                    st.dataframe(comp, use_container_width=True)
                    st.info("贡献度分析用于识别两个指定时间段之间，主指标变化主要由哪些维度项推动。正值代表拉动增长，负值代表拖累下降。")
                else:
                    st.info("请选择两个不同时间段进行贡献度分析。")
            else:
                st.info("当前时间数据不足以计算贡献度。")
        else:
            st.info("贡献度分析需要有效日期字段和至少一个维度字段。若无日期字段，可使用维度对比和结构分析替代。")

    with sub3:
        relation_candidates = [c for c in numeric_cols if not is_id_like(c)]
        if len(relation_candidates) < 2:
            st.warning("可用于关系分析的数值指标不足。系统会自动排除 ID/编号类字段。")
        else:
            c1, c2, c3 = st.columns([1, 1, 1.2])
            with c1:
                x_col = st.selectbox("X轴指标", relation_candidates, index=relation_candidates.index(main_metric) if main_metric in relation_candidates else 0)
            with c2:
                default_y_idx = 1 if len(relation_candidates) > 1 and relation_candidates[0] == x_col else 0
                y_col = st.selectbox("Y轴指标", relation_candidates, index=default_y_idx)
            with c3:
                group_col = st.selectbox("分组字段（可选）", ["无"] + selected_dimensions, index=1 if selected_dimensions else 0, key="relation_group")
                group_col = None if group_col == "无" else group_col
            rec_method, rec_reason = recommend_fit_method(df, x_col, y_col, group_col)
            methods = ["全局OLS（直线）", "按维度分组OLS（分段直线）", "LOWESS（局部加权）", "多项式回归（二阶）"]
            st.markdown(f'<div class="tip-card"><b>智能推荐：</b>{rec_method}。{rec_reason}</div>', unsafe_allow_html=True)
            fit_method = st.radio("选择拟合方法", methods, index=methods.index(rec_method), horizontal=True, key="fit_method_radio")
            corr = scatter_chart(df, x_col, y_col, group_col, fit_method)
            st.markdown(f'<div class="tip-card">{relationship_explanation(x_col, y_col, corr)}</div>', unsafe_allow_html=True)
            st.caption("全局OLS用于观察整体线性方向；分组OLS用于比较不同群体斜率；LOWESS用于发现局部拐点和非线性；二阶多项式用于观察先升后降或先降后升关系。")

            st.markdown("### 全部经营指标相关系数汇总")
            corr_matrix, corr_pairs = correlation_summary_tables(df, relation_candidates)
            if len(corr_pairs):
                st.markdown('<div class="highlight-note">相关系数表用于汇总所有经营数值指标之间的线性关系。数值越接近 1 表示正向关系越强，越接近 -1 表示反向关系越强，越接近 0 表示线性关系较弱。该结果用于发现线索，不直接等同于因果关系。</div>', unsafe_allow_html=True)
                st.dataframe(corr_pairs, use_container_width=True, hide_index=True, height=360)
                with st.expander("查看相关系数矩阵表"):
                    st.dataframe(corr_matrix, use_container_width=True)
            else:
                st.info("当前可用于相关系数汇总的有效数值指标不足。")

    with sub4:
        heat_type = st.radio("选择热力图类型", ["数值指标相关性矩阵", "维度交叉分布热力图"], horizontal=True)
        if heat_type == "数值指标相关性矩阵":
            corr_heatmap_chart(df, numeric_cols)
        else:
            if len(selected_dimensions) >= 2:
                row_dim = st.selectbox("行维度", selected_dimensions, key="heat_row")
                col_dim = st.selectbox("列维度", [d for d in selected_dimensions if d != row_dim], key="heat_col")
                agg_name = st.selectbox("聚合方式", ["求和", "均值", "计数"], index=0)
                topn = st.slider("TopN", 5, 20, 10)
                if agg_name == "计数":
                    row_top = df.groupby(row_dim).size().sort_values(ascending=False).head(topn).index
                    col_top = df.groupby(col_dim).size().sort_values(ascending=False).head(topn).index
                else:
                    row_top = df.groupby(row_dim)[main_metric].sum().sort_values(ascending=False).head(topn).index
                    col_top = df.groupby(col_dim)[main_metric].sum().sort_values(ascending=False).head(topn).index
                temp = df[df[row_dim].isin(row_top) & df[col_dim].isin(col_top)]
                if agg_name == "求和":
                    pivot = temp.pivot_table(index=row_dim, columns=col_dim, values=main_metric, aggfunc="sum", fill_value=0)
                elif agg_name == "均值":
                    pivot = temp.pivot_table(index=row_dim, columns=col_dim, values=main_metric, aggfunc="mean", fill_value=0)
                else:
                    pivot = temp.pivot_table(index=row_dim, columns=col_dim, values=main_metric, aggfunc="count", fill_value=0)
                heatmap_chart(pivot, f"{row_dim} × {col_dim} 的 {main_metric} 交叉分布热力图（{agg_name}）")
                st.info("交叉分布热力图用于观察两个维度组合下主指标的分布差异；数值指标相关性矩阵用于观察数值字段之间的相关关系。")
            else:
                st.info("维度交叉分布热力图需要至少两个维度字段。")


# ============================================================
# Tab 4
# ============================================================

with tab4:
    st.subheader("⚠️ 风险识别")
    st.markdown("""
    <div class="tip-card">
    <b>异常诊断说明：</b><br>
    本模块中的“异常”并不等同于数据错误，也不预设某一种固定经营场景。系统会先基于当前上传数据中实际存在的字段构建经营单元，再从四个方面综合判断风险：<br>
    ① <b>主指标分布偏离</b>：主指标处于当前样本的高位或低位区间；<br>
    ② <b>多指标组合偏离</b>：多个数值字段同时处于极端区间，说明该单元与整体样本差异较大；<br>
    ③ <b>动态业务规则</b>：仅当数据中真实存在利润、成本、费用、比例、预算等字段时，才生成对应的业务异常标签；<br>
    ④ <b>模型异常贡献</b>：Isolation Forest 根据多指标组合识别出的离群程度。<br>
    风险得分越高，表示该经营单元越需要优先核查；系统不会根据不存在的字段生成异常结论。
    </div>
    """, unsafe_allow_html=True)

    if len(anomaly_df) == 0:
        st.warning("当前数据不足以进行异常诊断。")
    else:
        abnormal = anomaly_df[anomaly_df["是否经营异常"]].copy()
        high = abnormal[abnormal["风险等级"] == "高风险"]

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            kpi_card("异常经营单元", f"{len(abnormal)}", "中高风险单元数量")
        with c2:
            kpi_card("高风险单元", f"{len(high)}", "风险得分≥70")
        with c3:
            kpi_card("最高风险得分", f"{anomaly_df['风险得分'].max():.1f}", "越高表示偏离越明显")
        with c4:
            kpi_card("异常占比", pct_fmt(len(abnormal) / max(len(anomaly_df), 1)), "异常单元 / 全部单元")

        st.markdown("### 风险地图")
        c_map1, c_map2 = st.columns([1.35, 0.9])
        with c_map1:
            risk_matrix_chart(anomaly_df, main_metric)
        with c_map2:
            risk_source_summary_cards(anomaly_df)

        c1, c2 = st.columns(2)
        with c1:
            fig = go.Figure(go.Histogram(x=anomaly_df["风险得分"], nbinsx=20))
            fig.update_layout(title="风险得分分布", xaxis_title="风险得分", yaxis_title="数量")
            st.plotly_chart(chart_layout(fig, 360), use_container_width=True)
        with c2:
            rc = anomaly_df["风险等级"].value_counts().reset_index()
            rc.columns = ["风险等级", "数量"]
            bar_chart(rc, "风险等级", "数量", "风险等级分布")

        st.markdown("### Top 异常经营单元")
        st.markdown("""
        <div class="highlight-note">
        系统根据字段业务属性自动匹配分析口径：金额、数量、利润等规模型指标采用汇总展示，折扣、比率、单价、时长等水平型指标采用平均水平展示，使结果更贴近经营管理语境。
        </div>
        """, unsafe_allow_html=True)
        top_anom = anomaly_df.sort_values("风险得分", ascending=False).head(10)
        st.dataframe(top_anom, use_container_width=True, height=320)

        st.markdown("### 异常解释与风险组成")
        if len(top_anom):
            chosen_idx = st.selectbox("选择一个异常单元查看解释", top_anom.index.tolist())
            row = anomaly_df.loc[chosen_idx]

            ignore_cols = set(numeric_cols + ["记录数", "是否AI异常", "异常得分", "异常依据", "主指标偏离", "多指标组合偏离", "业务规则风险", "模型异常贡献", "风险得分", "风险等级", "是否经营异常"])
            title_cols = [c for c in anomaly_df.columns if c not in ignore_cols]
            obj = "｜".join([f"{c}={row[c]}" for c in title_cols[:3]]) if title_cols else f"记录 {chosen_idx}"

            c1, c2 = st.columns([1, 1])
            with c1:
                st.markdown(f"""
                <div class="section-card">
                <h4>为什么是异常？</h4>
                <p class="small-text"><b>对象：</b>{obj}</p>
                <p class="small-text"><b>风险等级：</b>{row['风险等级']}｜<b>风险得分：</b>{row['风险得分']}</p>
                <p class="small-text"><b>异常依据：</b>{row['异常依据']}</p>
                </div>
                """, unsafe_allow_html=True)
            with c2:
                comp = pd.DataFrame({
                    "风险来源": ["主指标偏离", "多指标组合偏离", "业务规则风险", "模型异常贡献"],
                    "得分": [row["主指标偏离"], row["多指标组合偏离"], row["业务规则风险"], row["模型异常贡献"]]
                })
                radar_chart(comp)

            st.markdown("#### 风险得分构成")
            st.dataframe(comp, use_container_width=True)

            with st.expander("🤖 使用大模型生成个性化异常原因与行动建议"):
                provider = st.selectbox("选择大模型", list(get_llm_configs().keys()), key="anom_llm")
                if st.button("生成AI异常解释", key="anom_ai_btn"):
                    try:
                        sys_p = "你是经营风险诊断专家。请根据异常单元数据，解释为什么异常、为什么高风险，并给出可执行建议。"
                        user_p = f"异常对象：{obj}\n风险等级：{row['风险等级']}\n风险得分：{row['风险得分']}\n异常依据：{row['异常依据']}\n风险得分构成：{comp.to_dict(orient='records')}\n该单元完整数据：{row.to_dict()}\n\n请输出：1. 异常原因；2. 高风险判断依据；3. 可能影响；4. 管理建议；5. 下一步核查数据。"
                        ans, elapsed = call_llm(provider, sys_p, user_p, temperature=0.2)
                        st.success(f"生成完成，耗时 {elapsed:.2f} 秒")
                        render_ai_answer_pretty(ans)
                    except Exception as e:
                        st.error(f"AI异常解释生成失败：{e}")

        st.markdown("### 异常依据补充")
        st.info("上表中的“异常依据”会说明具体触发原因，例如主指标处于高位区间、多项指标同时极端、或实际存在字段触发业务规则。建议结合原始明细进一步核查。")


# ============================================================
# Tab 5
# ============================================================

with tab5:
    st.subheader("💬 问数助手")
    st.markdown("""
    <div class="tip-card">
    用户可以输入与当前数据相关的问题。系统会调用所选大模型生成 SQL，并在数据库中执行查询；同时展示
    <b>AI直接回答、问题解析口径、SQL语句、执行结果、系统校验和评分明细</b>，用于降低字段误解和模型幻觉风险。
    </div>
    """, unsafe_allow_html=True)

    query_process_visual()

    with st.expander("查看问数任务综合评分规则"):
        st.markdown("""
        **问数任务综合评分**用于比较不同大模型在本轮数据问数任务中的表现，不等同于严格人工标注准确率。评分规则如下：

        1. **SQL可执行性（40分）**：生成的SQL必须是安全的 `SELECT` 查询，并且能在当前上传数据表上成功运行；若字段名错误、语法错误或生成非查询语句，则该项为0分。  
        2. **查询结果有效性（20分）**：SQL成功执行且返回非空结果，通常给20分；如果成功执行但结果为空，说明查询条件或理解可能有偏差，通常只给8分。  
        3. **字段匹配度（15分）**：系统会检查SQL是否使用了当前数据中的真实字段。使用了相关字段通常给15分；虽然可执行但字段使用较弱时会降分。  
        4. **结果解释质量（15分）**：模型需要基于查询结果给出清晰的经营解释，包括结论、含义和建议；解释过短或缺失会扣分。  
        5. **响应速度（10分）**：响应时间≤3秒为10分；3—8秒为8分；8—15秒为5分；超过15秒为2分；调用失败为0分。  

        分数越高表示本轮问数结果越“可用”。但由于自然语言问题可能存在口径歧义，用户仍应结合“问题解析口径”和“SQL查询结果”进行复核。
        """)

    st.markdown("### 💡 你可以这样问")
    examples = generate_dynamic_questions(main_metric, selected_dimensions, numeric_cols, valid_date_col)
    rows = [examples[i:i+3] for i in range(0, len(examples), 3)]
    for row_qs in rows[:2]:
        cols = st.columns(3)
        for col_obj, (typ, q) in zip(cols, row_qs):
            with col_obj:
                st.markdown(f"""
                <div class="example-card">
                    <div class="example-title">{typ}</div>
                    <div class="example-q">{q}</div>
                </div>
                """, unsafe_allow_html=True)

    providers = list(get_llm_configs().keys())
    selected_providers = st.multiselect("选择参与问数/对比的大模型", providers, default=providers)
    question = st.text_input("请输入问题", value=examples[0][1] if examples else f"请分析{main_metric}的整体情况")

    if "llm_records" not in st.session_state:
        st.session_state["llm_records"] = []

    if st.button("开始问数", type="primary"):
        if not selected_providers:
            st.warning("请至少选择一个大模型。")
        else:
            for provider in selected_providers:
                with st.spinner(f"{provider} 正在生成 SQL、执行查询并生成可信度说明..."):
                    try:
                        pack = run_llm_sql_question(provider, question, df, meta)
                    except Exception as e:
                        pack = {"模型": provider, "SQL": "", "是否成功": False, "错误信息": str(e), "结果": pd.DataFrame(), "解释": "", "直接回答": "", "响应时间": np.nan, "综合评分": 0, "评分明细": {}}
                        pack["问题解析口径"] = infer_question_caliber(question, "", pd.DataFrame(), df, meta)
                        pack["系统校验"] = build_query_validation(pack, question, df)

                render_llm_query_pack(pack, question, df, meta)

                if pack.get("是否成功"):
                    st.markdown("#### 推荐追问")
                    followups = recommend_followup_questions(question, main_metric, selected_dimensions, numeric_cols)
                    st.markdown(" ".join([f'<span class="governance-step">{q}</span>' for q in followups]), unsafe_allow_html=True)

                st.session_state["llm_records"].append({
                    "模型": provider,
                    "问题": question,
                    "SQL可执行": "是" if pack["是否成功"] else "否",
                    "结果行数": len(pack["结果"]) if isinstance(pack["结果"], pd.DataFrame) else 0,
                    "响应时间秒": round(pack["响应时间"], 2) if not pd.isna(pack["响应时间"]) else None,
                    "问数任务综合评分": pack["综合评分"],
                    "错误信息": pack["错误信息"]
                })

    st.markdown("### 模型本轮/历史对比")
    records = st.session_state.get("llm_records", [])
    if records:
        compare_df = pd.DataFrame(records)
        st.dataframe(compare_df.tail(20), use_container_width=True)
        summary = compare_df.groupby("模型").agg(
            测试次数=("模型", "count"),
            成功次数=("SQL可执行", lambda x: (x == "是").sum()),
            平均响应时间=("响应时间秒", "mean"),
            平均评分=("问数任务综合评分", "mean")
        ).reset_index()
        bar_chart(summary, "模型", "平均评分", "平均问数任务综合评分")
    else:
        st.info("暂无问数记录。")


# ============================================================
# Tab 6
# ============================================================

with tab6:
    st.subheader("📝 决策简报生成")

    focus = st.multiselect(
        "选择本次简报关注重点",
        ["综合经营", "趋势变化", "维度结构", "异常风险", "成本/费用", "利润/收益", "预算执行", "人力/薪酬", "采购/供应商"],
        default=["综合经营", "维度结构", "异常风险"]
    )

    st.markdown("### 决策简报预览")
    st.info(f"系统将基于当前上传数据、主分析指标“{main_metric}”和用户选择的关注重点生成 Word 决策简报。下方展示的是即将写入报告的核心分析口径和代表性图表，完整报告将在点击生成后下载。当前数据共 {len(df):,} 条记录。")

    if selected_dimensions:
        dim = selected_dimensions[0]
        full_g = dimension_summary(df, dim, main_metric)
        g = full_g.head(10)
        chart_title = f"代表性图表预览：按{dim}汇总{main_metric}" if len(full_g) <= 10 else f"代表性图表预览：按{dim}汇总{main_metric} Top10"
        bar_chart(g, dim, f"{main_metric}合计", chart_title)

    if st.button("生成并下载 Word 简报", type="primary"):
        try:
            bio = generate_report_docx(df, main_metric, selected_dimensions, valid_date_col, anomaly_df, focus)
            st.download_button(
                "下载 Word 简报",
                data=bio,
                file_name="智策经营_管理层决策简报.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Word 简报生成失败：{e}")

    with st.expander("使用大模型生成更丰富的简报正文"):
        provider = st.selectbox("选择大模型", list(get_llm_configs().keys()), key="report_llm")
        if st.button("生成AI简报正文", key="report_ai_btn"):
            try:
                abnormal_count = int(anomaly_df["是否经营异常"].sum()) if len(anomaly_df) else 0
                top_anom = anomaly_df.sort_values("风险得分", ascending=False).head(5).to_dict(orient="records") if len(anomaly_df) else []
                sys_p = "你是企业经营分析顾问。请站在管理层阅读者角度，基于系统统计结果生成补充解读和可执行管理建议。不得引用数据中不存在的字段，不要输出报告标题，不要重新生成一、二、三、四、五大纲。"
                user_p = f"主指标：{main_metric}\n关注重点：{'、'.join(focus)}\n数据记录数：{len(df)}\n主指标合计：{df[main_metric].sum()}\n主指标均值：{df[main_metric].mean()}\n维度字段：{selected_dimensions}\n异常经营单元数：{abnormal_count}\nTop异常：{top_anom}\n\n请输出两部分：1）AI综合判断：用2-4段话说明当前数据反映的主要经营问题和管理含义；2）管理行动清单：用Markdown表格输出“管理关注点、需要解决的问题、建议动作、优先级”。重点说明为什么这些异常值得关注、应该查哪些原始数据、如何形成下一步管理动作。"
                ans, elapsed = call_llm(provider, sys_p, user_p, temperature=0.25)
                st.success(f"生成完成，耗时 {elapsed:.2f} 秒")
                render_ai_answer_pretty(ans)
                try:
                    ai_bio = generate_ai_report_docx(ans, df, main_metric, selected_dimensions, valid_date_col, anomaly_df, focus)
                    st.download_button(
                        "下载 AI增强版 Word 简报",
                        data=ai_bio,
                        file_name="智策经营_AI增强版决策简报.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as doc_e:
                    st.warning(f"AI简报正文已生成，但Word下载文件生成失败：{doc_e}")
            except Exception as e:
                st.error(f"AI简报正文生成失败：{e}")
